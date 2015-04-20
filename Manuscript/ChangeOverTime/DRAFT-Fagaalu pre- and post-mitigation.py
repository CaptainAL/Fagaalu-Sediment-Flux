# -*- coding: utf-8 -*-
"""
Created on Thu Mar 19 12:14:25 2015

@author: Alex
"""

# -*- coding: utf-8 -*-
"""
Created on Tue Feb 17 10:50:00 2015

@author: Alex
"""
plt.ioff()
plt.close('all')
from docx import *
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
##  Create Document
document = Document()

######## SOME TOOLS
def add_figure_caption(fig_num=str(len(document.inline_shapes)),caption=''):
    cap = document.add_paragraph("Figure "+fig_num+". "+caption)
    cap.paragraph_style = 'caption'
    return
    
def dataframe_to_table(df=pd.DataFrame(),table_num=str(len(document.tables)+1),caption='',fontsize=11):
    table = document.add_table(rows=1, cols=len(df.columns)) 
    ## Merge all cells in top row and add caption text
    table_caption = table.rows[0].cells[0].merge(table.rows[0].cells[len(df.columns)-1])
    table_caption.text = "Table "+table_num+". "+caption
    ## Add  header
    header_row = table.add_row().cells
    col_count =0 ## counter  to iterate over columns
    for col in  header_row:
        col.text = df.columns[col_count] #df.columns[0] is the index
        col_count+=1
    ## Add data by  iterating over the DataFrame rows, then using a dictionary of DataFrame column labels to extract data
    col_labels = dict(zip(range(len(df.columns)),df.columns.tolist())) ## create dictionary where '1  to  n' is key for DataFrame columns
    for row in df.iterrows():  ## iterate over  rows in  DataFrame
        #print row[1]
        row_cells = table.add_row().cells ## Add a row to the  table
        col_count  = 0
        for cell in row_cells: ## iterate over the columns in the row
            #print row[1][str(col_labels[col_count])]
            cell.text = str(row[1][str(col_labels[col_count])]) ## and plug in data using a dictionary to  get column labels for DataFrame
            col_count+=1
    ## Format Table Style  
    table.style = 'TableGrid' 
    table.style.font.size = Pt(fontsize)
    table.autofit
    #table.num = str(len(document.tables)+1)
    return table
    
def add_equation(eq_table):
    t = document.add_table(rows=len(eq_table.rows),cols=len(eq_table.columns))
    t.rows[0].cells[1].text = eq_table.rows[0].cells[1].text
    t.rows[0].cells[2].text = eq_table.rows[0].cells[2].text   
    if len(eq_table.rows)>1:
        t.rows[1].cells[0].merge(t.rows[1].cells[2]).text = eq_table.rows[1].cells[0].text
    t.style = 'TableGrid'   
    t.autofit
    return t
###################################################################################################################################################################    


#### TABLES ########################################################################################################################################################
### Landcover_Table
table_count=0
def tab_count():
    global table_count
    table_count+=1
    return str(table_count)
# Prepare LULC Data
landcover_table = LandCover_table()
landcover_table.table_num = str(tab_count())

### Storm Sediment Discharge Table
## Prepare table data
S_Diff_table = S_storm_diff_table(subset='post')[0] ## function to create table data
S_Diff_table.table_num = str(tab_count())

S_Diff_table_quarry = S_storm_diff_table_quarry(subset='post',manual_edit=False)[0]
S_Diff_table_quarry.table_num = str(tab_count())

### Storm Q and SSY summary table
#Q_S_Diff_summary_table = Q_S_storm_diff_summary_table(subset='post')
#Q_S_Diff_summary_table.table_num = str(tab_count())

### Model statistics table
SSYEV_models_stats = ALLRatings_table(subset='post')
SSYEV_models_stats.table_num = str(tab_count())

#### Appendix
table_count,figure_count,equation_count=0, 0, 0
### Storm Water Discharge Table
#Q_Diff_table = Q_storm_diff_table() ## function to create table data
#Q_Diff_table.table_num =str(tab_count())

#### FIGURES ########################################################################################################################################################
figure_count=0
def fig_count():
    global figure_count
    figure_count+=1
    return str(figure_count)
## INTRODUCTION
#### Study Area Map
Study_Area_map = {'filename':maindir+'Figures/Maps/FagaaluInstruments land only map.tif', 'fig_num':str(fig_count())}
## Quarry_picture
Quarry_picture_1 = {'filename':maindir+'Figures/Maps/Quarry_Mitigation/3_Retention ponds aerial.jpg','fig_num':str(fig_count())}
Quarry_picture_2 = {'filename':maindir+'Figures/Maps/Quarry_Mitigation/4_Retention ponds.jpg','fig_num':str(fig_count())}


## RESULTS

####  SSC
## SSC Boxplots
SSC_Boxplot_pre_and_posts= {'filename':figdir+'SSC/Grab sample boxplots baseflow and stormflow pre and post mitigation','fig_num':str(fig_count())}
plotSSCboxplots_pre_and_post(withR2=False,log=True,show=False,save=True,filename=SSC_Boxplots['filename'])

## Discharge vs Sediment Concentration
Discharge_Concentration = {'filename':figdir+'SSC/Water discharge vs Sediment concentration pre and post mitigation','fig_num':str(fig_count())}
plotQvsC_pre_and_post(storm_samples_only=False,ms=6,show=False,log=False,save=True,filename=Discharge_Concentration['filename'])

####  T-SSC Rating Curves
## Synthetic Rating Curves
Synthetic_Rating_Curve = {'filename':figdir+'T/Synthetic Rating Curves Fagaalu','fig_num':str(fig_count())} ## define file name to find the png file from other script
Synthetic_Rating_Curves_Fagaalu(param='SS_Mean',show=False,save=True,filename=Synthetic_Rating_Curve['filename'])## generate figure from separate script and save to file
## LBJ and DAM YSI T-SSC rating curves
LBJ_and_DAM_YSI_Rating_Curve = {'filename':figdir+'T/T-SSC rating LBJ and DAM YSI','fig_num':str(fig_count())}
plotYSI_compare_ratings(DAM_YSI,DAM_SRC,LBJ_YSI,Use_All_SSC=False,show=False,save=True,filename=LBJ_and_DAM_YSI_Rating_Curve['filename'])
## LBJ OBSa T-SSC rating curve
LBJ_OBSa_Rating_Curve = {'filename':figdir+'T/T-SSC rating LJB OBSa','fig_num':str(fig_count())}
OBSa_compare_ratings(df=LBJ_OBSa,df_SRC=LBJ_SRC,SSC_loc='LBJ',Use_All_SSC=False,show=False,save=True,filename=LBJ_OBSa_Rating_Curve['filename'])  
## LBJ OBSa T-SSC rating curve
LBJ_OBSb_Rating_Curve = {'filename':figdir+'T/T-SSC rating LJB OBSb','fig_num':str(fig_count())}
OBSb_compare_ratings(df=LBJ_OBSb,df_SRC=LBJ_SRC,SSC_loc='LBJ',Use_All_SSC=False,show=False,save=True,filename=LBJ_OBSb_Rating_Curve['filename'])  

#### Storms
Pre_Example_Storm = {'filename':figdir+'storm_figures/Pre_Example_Storm','fig_num':str(fig_count())}
plot_storm_individually(LBJ_storm_threshold,LBJ_StormIntervals.loc[63],show=False,save=True,filename=Pre_Example_Storm['filename']) 

Post_Example_Storm = {'filename':figdir+'storm_figures/Post_Example_Storm','fig_num':str(fig_count())}
plot_storm_individually(LBJ_storm_threshold,LBJ_StormIntervals.loc[116],show=False,save=True,filename=Post_Example_Storm['filename']) 

#### SSY models
SSY_models_ALL = {'filename':figdir+'SSY/SSY Models ALL post mitigation','fig_num':str(fig_count())}
ALLStorms_ALLRatings = plotALLStorms_ALLRatings(subset='post',ms=4,norm=True,log=True,show=False,save=True,filename=SSY_models_ALL['filename'])

###### EQUATIONS ############################################################################################################################################
equation_count=0
def eq_count():
    global equation_count
    equation_count+=1
    return str(equation_count)
    
Equations = Document(maindir+'/Manuscript/Equations.docx').tables
## SSYev = Q*SSC 
SSYEV_eq = Equations[0].table
SSYEV_eq.eq_num = eq_count()
## DR = SSY/SSYPRE
DR_eq = Equations[2].table
DR_eq.eq_num = eq_count()
##  SSYPRE = Area_total * (SSYupper/AREAupper)
SSYPRE_eq = Equations[3].table
SSYPRE_eq.eq_num = eq_count()
## predict_SSYev = aXb
predict_SSYEV_eq = Equations[1].table
predict_SSYEV_eq.eq_num = eq_count()
## PE = sqrt(sum(Error^2+Error^2))
PE_eq = Equations[4].table
PE_eq.eq_num = eq_count()
############################################################################################################################################

#### TITLE
title_title = document.add_heading('TITLE:',level=1)
title_title.paragraph_format.space_before = 0
title = document.add_heading('Contributions of human activities to suspended sediment yield during storm events from a steep, small, tropical watershed',level=1)
title.paragraph_format.space_before = 0
## subscript/superscript words
document.add_paragraph("SSYEV, m3, km2, SSYUPPER, SSYLOWER, SSYTOTAL, SSYPRE, SSYFOREST, SSYVILLAGE, SSYQUARRY, alpha, Beta")

#### ABSTRACT
abstract_title = document.add_heading('ABSTRACT',level=2)
abstract_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract = document.add_paragraph('Abstract text goes here....')
  
#### INTRODUCTION
introduction_title = document.add_heading('Introduction',level=2)
#### STUDY AREA
study_area_title = document.add_heading('Study Area',level=2)
## Study Area map
if 'Study_Area_map' in locals():
    document.add_picture(Study_Area_map['filename'],width=Inches(6))
    add_figure_caption(Study_Area_map['fig_num'],"Faga'alu watershed showing the Uupper (undisturbed) and Lower (human-disturbed) subwatersheds.")

#### METHODS
methods_title = document.add_heading('Methods',level=2)

#### RESULTS ####
results_title = document.add_heading('Results and Discussion',level=2)
## Field data collection
document.add_heading('Field Data Collection',level=3)


#### Suspended Sediment Concentration
document.add_heading('Suspended Sediment Concentration',level=4)
def No_All_samples(location,subset='Post-ALL'):
    No_samples = len(SSC_dict[subset][SSC_dict[subset]['Location'].isin([location])])
    return No_samples
## Mean and Max SSC  numbers
def Mean_and_Max_SSC(ssc,location):
    mean =  int(ssc[ssc['Location'].isin([location])]['SSC (mg/L)'].mean())
    maximum = int(ssc[ssc['Location'].isin([location])]['SSC (mg/L)'].max())
    return "{:,}".format(mean), "{:,}".format(maximum)
Mean_SSC_Forest, Max_SSC_Forest  = Mean_and_Max_SSC(SSC_dict['Post-ALL'],'DAM')
Mean_SSC_Quarry, Max_SSC_Quarry = Mean_and_Max_SSC(SSC_dict['Post-ALL'],'DT')
Mean_SSC_Village, Max_SSC_Village = Mean_and_Max_SSC(SSC_dict['Post-ALL'],'LBJ')
Max_SSC_Quarry_Event = "{:%m/%d/%Y}".format(SSC_dict['Post-ALL'][SSC_dict['Post-ALL']['Location'].isin(['DT'])]['SSC (mg/L)'].idxmax())
Max_SSC_Quarry_Q = "{:,}".format(int(QUARRY['Q'][SSC_dict['Post-ALL'][SSC_dict['Post-ALL']['Location'].isin(['DT'])]['SSC (mg/L)'].idxmax()]))
Max_SSC_Village_Event = "{:%m/%d/%Y}".format(SSC_dict['Post-ALL'][SSC_dict['Post-ALL']['Location'].isin(['LBJ'])]['SSC (mg/L)'].idxmax())
Max_SSC_Village_Q = "{:,}".format(int(LBJ['Q'][SSC_dict['Post-ALL'][SSC_dict['Post-ALL']['Location'].isin(['LBJ'])]['SSC (mg/L)'].idxmax()]))
Max_SSC_Forest_Event = "{:%m/%d/%Y}".format(SSC_dict['Post-ALL'][SSC_dict['Post-ALL']['Location'].isin(['DAM'])]['SSC (mg/L)'].idxmax())
Max_SSC_Forest_Q = "{:,}".format(int(DAM['Q'][SSC_dict['Post-ALL'][SSC_dict['Post-ALL']['Location'].isin(['DAM'])]['SSC (mg/L)'].idxmax()]))

document.add_paragraph("From October 1, 2014, to January 9, 2015, "+str(len(SSC_dict['Post-ALL']))+" samples were collected in Faga'alu and analyzed for SSC. Three sites were the focus of this analysis: 1) FOREST(n="+"%.0f"%No_All_samples('DAM')+"), 2) QUARRY (n="+"%.0f"%No_All_samples('DT')+" grab samples, n="+"%.0f"%No_All_samples('R2')+" from the Autosampler), and 3) VILLAGE (n="+"%.0f"%No_All_samples('LBJ')+"). Mean SSC of grab samples collected during baseflow and stormflow were "+Mean_SSC_Forest+" mg/L at the FOREST site, "+Mean_SSC_Quarry+" mg/L at the QUARRY, and "+Mean_SSC_Village+" mg/L the downstream VILLAGE site. Maximum SSC values were "+Max_SSC_Forest+" mg/L, "+Max_SSC_Quarry+" mg/L, and "+Max_SSC_Village+" mg/L at the upstream (FOREST), quarry (QUARRY), and downstream (VILLAGE) sites, respectively. The maximum SSC value at QUARRY ("+Max_SSC_Quarry+" mg/L ) was sampled on "+Max_SSC_Village_Event+" at "+Max_SSC_Quarry_Q+" L/sec. The maximum SSC value at VILLAGE ("+Max_SSC_Village+" mg/L) was sampled on "+Max_SSC_Village_Event+", when discharge was "+Max_SSC_Village_Q+" L/sec. The maximum SSC value for the upstream site ("+Max_SSC_Forest+" mg/L) was sampled on "+Max_SSC_Forest_Event+" at discharge "+Max_SSC_Forest_Q+" L/sec.")

def No_storm_samples(location):
    No_storm = len(SSC_dict['Post-storm'][SSC_dict['Post-storm']['Location'].isin([location])])
    return "%.0f"%No_storm
# Percent storm samples in another script, wouldn't work here
def Mean_storm_SSC(location):    
    Mean_storm_samples  = SSC_dict['Post-storm'][SSC_dict['Post-storm']['Location'].isin([location])]['SSC (mg/L)'].mean()
    return "%.0f"%Mean_storm_samples 
def Percent_storm_samples(location):
    storm_samples = len(SSC_dict['Post-storm'][SSC_dict['Post-storm']['Location'].isin([location])])
    all_samples = len(SSC_dict['Post-ALL'][SSC_dict['Post-ALL']['Location'].isin([location])]) 
    percent_storm = storm_samples/all_samples *100
    return "%.0f"%percent_storm
Percent_storm_Forest = Percent_storm_samples('DAM')
Percent_storm_Quarry = Percent_storm_samples('DT')
Percent_storm_Village = Percent_storm_samples('LBJ')

DAM_Stormflow_conditions = DAM[DAM['stage']==DAM_storm_threshold.round(0)]['Q'][0]
LBJ_Stormflow_conditions = LBJ[LBJ['stage']==LBJ_storm_threshold.round(0)]['Q'][0]

document.add_paragraph("At FOREST, "+No_storm_samples('DAM')+" grab samples ("+Percent_storm_Forest+"%) were taken during stormflow conditions (Q_DAM>"+"%.0f"%DAM_Stormflow_conditions+" L/sec), with mean SSC of "+Mean_storm_SSC('DAM')+" mg/L. At QUARRY, "+No_storm_samples('DT')+" grab samples ("+Percent_storm_Quarry +"%) were taken during stormflow conditions (Q_DAM>"+"%.0f"%DAM_Stormflow_conditions+" L/sec), with mean SSC of "+Mean_storm_SSC('DT')+" mg/L. At VILLAGE, "+No_storm_samples('LBJ')+" samples ("+Percent_storm_Village+"%) were taken during stormflow conditions (Q_VILLAGE>"+"%.0f"%LBJ_Stormflow_conditions+" L/sec), with mean SSC of "+Mean_storm_SSC('LBJ')+" mg/L (Figure "+SSC_Boxplots['fig_num']+"). This pattern of SSC values suggests that little sediment is contributed from the forest upstream of FOREST, then there is a large input of sediment between FOREST and QUARRY, and then SSC is diluted by addition of stormflow with lower SSC between QUARRY and VILLAGE.")
## SSC boxplots
if 'SSC_Boxplots' in locals():
    document.add_picture(SSC_Boxplots['filename']+'.png',width=Inches(6))
    add_figure_caption(SSC_Boxplots['fig_num'],"Boxplots of Suspended Sediment Concentration (SSC) from grab samples only (no Autosampler) at FOREST, QUARRY, and VILLAGE during storm periods.")
## Water Discharge vs Sediment Concentration
if 'Discharge_Concentration' in locals():
    document.add_picture(Discharge_Concentration['filename']+'.png',width=Inches(6))
    add_figure_caption(Discharge_Concentration['fig_num'],"Water Discharge vs Suspended Sediment Concentration at FOREST, QUARRY, and VILLAGE during baseflow and stormflow periods. Samples from Autosampler are included with grab samples at QUARRY.")

#### Turbidity
document.add_heading('Turbidity', level=4)
document.add_paragraph("A "+'"synthetic"'+" T-SSC relationship (SRC) was developed for the YSI 600OMS turbidimeter at FOREST and the CampbellSci OBS500 at VILLAGE, in 2014, using sediment collected from the streambed near the turbidimeter (Figure "+Synthetic_Rating_Curve['fig_num']+"). These relationships were compared with the T-SSC relationships developed using in situ data collected under storm conditions (Figures "+LBJ_and_DAM_YSI_Rating_Curve['fig_num']+"-"+LBJ_OBSb_Rating_Curve['fig_num']+"). In all instances, the SRC's had a steeper slope than the T-SSC relationships from storm samples, suggesting that even though the sediment used to develop the SRC's was taken from nearby stream banks, it differed in character from the sediment collected under actual storm conditions. The T-SSC relationships under actual storm conditions all showed acceptable r2 values ("+"%.2f"%LBJ_YSI_rating.r2+"-"+"%.2f"%DAM_YSI_rating.r2+"), so the SRC's were not used to model SSC from T data.")
## Synthetic Rating Curves
if 'Synthetic_Rating_Curve' in locals():
    document.add_picture(Synthetic_Rating_Curve['filename']+'.png',width=Inches(6)) ## add pic from filename defined above
    add_figure_caption(Synthetic_Rating_Curve['fig_num'],"Synthetic Rating Curves for Turbidimeters deployed at FOREST and VILLAGE")

document.add_paragraph("The T-SSC relationship varied among sampling sites and sensors. Lower scatter in the linear relationships was achieved by using grab samples collected during stormflows only. It is assumed that the color, particle sizes, and composition of sediment changes during stormflows when sediment from the quarry, which is lighter in color and finer, is present. For the TS3000 deployed at FOREST, the r2 value was fairly high ("+"%.2f"%DAM_TS3K_rating.r2+") but the ranges of T and SSC values used to develop the relationship were considered too small to develop a robust relationship for higher T values. Instead, the T-SSC relationship developed for the YSI turbidimeter installed at FOREST (Figure "+LBJ_and_DAM_YSI_Rating_Curve['fig_num']+") was used to convert T data from the TS3000 to SSC. For the YSI 600OMS turbidimeter, more scatter was observed in the T-SSC relationship at VILLAGE than at FOREST (Figure "+LBJ_and_DAM_YSI_Rating_Curve['fig_num']+"), but this could be attributed to the higher number and wider range of values sampled, as well the contribution of multiple sediment sources sampled at VILLAGE.") 

## LBJ and DAM YSI T-SSC rating curves
if 'LBJ_and_DAM_YSI_Rating_Curve' in locals():
    document.add_picture(LBJ_and_DAM_YSI_Rating_Curve['filename']+'.png',width=Inches(6))
    add_figure_caption(LBJ_and_DAM_YSI_Rating_Curve['fig_num'],"Turbidity-Suspended Sediment Concentration relationships for the YSI turbidimeter deployed at VILLAGE ("+"{:%m/%d/%Y}".format(LBJ_YSI.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(LBJ_YSI.dropna().index[-1])+") and at FOREST ("+"{:%m/%d/%Y}".format(DAM_YSI.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(DAM_YSI.dropna().index[-1])+").")
    
document.add_paragraph("The CampbellSci OBS500 measures both BS and SS, and both the BS-SSC and SS-SSC relationships at VILLAGE showed high r2 values (Figure "+LBJ_OBSa_Rating_Curve['fig_num']+" and Figure "+LBJ_OBSb_Rating_Curve['fig_num']+"). Mean SS was used since it is physically measured the same as NTU measured by the YSI turbidimeter (Chauncey W. Anderson 2005).") 

## LBJ OBSa T-SSC rating curve
if 'LBJ_OBSa_Rating_Curve' in locals():
    document.add_picture(LBJ_OBSa_Rating_Curve['filename']+'.png',width=Inches(6))
    add_figure_caption(LBJ_OBSa_Rating_Curve['fig_num'],"Turbidity-Suspended Sediment Concentration relationships for the OBS500 turbidimeter deployed at VILLAGE ("+"{:%m/%d/%Y}".format(LBJ_OBSa.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(LBJ_OBSa.dropna().index[-1])+").")

## LBJ OBSb T-SSC rating curve
if 'LBJ_OBSb_Rating_Curve' in locals():
    document.add_picture(LBJ_OBSb_Rating_Curve['filename']+'.png',width=Inches(6))
    add_figure_caption(LBJ_OBSb_Rating_Curve['fig_num'],"Turbidity-Suspended Sediment Concentration relationships for the OBS500 turbidimeter deployed at VILLAGE ("+"{:%m/%d/%Y}".format(LBJ_OBSb.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(LBJ_OBSb.dropna().index[-1])+").")

document.add_paragraph("The RMSE for each selected T-SSC relationship was computed and used in the estimate of Probable Error. The RMSE for the T-SSC relationship was "+"%.0f"%DAM_YSI_rating.rmse+", "+"%.0f"%LBJ_YSI_rating.rmse+", "+"%.0f"%LBJ_OBSa_rating.rmse+", "+"%.0f"%LBJ_OBSb_rating.rmse+" mg/L for the YSI at FOREST, YSI at VILLAGE, CampbellSci OBS500 at VILLAGE (2013), and CampbellSci OBS500 at VILLAGE (2014), respectively. The RMSE for YSI at FOREST was also used for the TS3000 at FOREST since the same T-SSC relationship was used.")

#### Storm Events
document.add_heading('Storm Events',level=3)
No_of_Storm_Intervals = len(LBJ_StormIntervals[LBJ_StormIntervals['start']>Mitigation])
No_of_Storm_Intervals_DAM_Q = len(SedFluxStorms_DAM[SedFluxStorms_DAM['Qstart']>Mitigation]['Qsum'].dropna())
No_of_Storm_Intervals_LBJ_Q = len(SedFluxStorms_LBJ[SedFluxStorms_LBJ['Qstart']>Mitigation]['Qsum'].dropna())
No_of_Storm_Intervals_DAM_S = len(SedFluxStorms_DAM[SedFluxStorms_DAM['Sstart']>Mitigation]['Ssum'].dropna())
No_of_Storm_Intervals_LBJ_S = len(SedFluxStorms_LBJ[SedFluxStorms_LBJ['Sstart']>Mitigation]['Ssum'].dropna())
No_of_Storm_Intervals_QUA_S = len(S_storm_diff_table_quarry(subset='post',manual_edit=False))-1
max_storm_duration = LBJ_StormIntervals[LBJ_StormIntervals['start']>Mitigation]['duration (hrs)'].max()/24


if 'Pre_Example_Storm' in locals():
    document.add_picture(Pre_Example_Storm['filename']+'.png',width=Inches(6))
    add_figure_caption(Pre_Example_Storm['fig_num'],"Example of storm event ("+"{:%m/%d/%Y}".format(LBJ_StormIntervals.loc[63]['start'])+") before sediment mitigation structures installed at quarry. SSY at FOREST and VILLAGE calculated from SSC modeled from T, and SSY at QUARRY from SSC samples collected by the Autosampler.")
    
if 'Post_Example_Storm' in locals():
    document.add_picture(Post_Example_Storm['filename']+'.png',width=Inches(6))
    add_figure_caption(Post_Example_Storm['fig_num'],"Example of storm event ("+"{:%m/%d/%Y}".format(LBJ_StormIntervals.loc[116]['start'])+") after sediment mitigation structures installed at quarry. SSY at FOREST and VILLAGE calculated from SSC modeled from T, and SSY at QUARRY from SSC samples collected by the Autosampler.")
    
document.add_paragraph("Using the stage threshold method and manual separation of complex storm events, "+"%.0f"%No_of_Storm_Intervals+" storm events were identified from Q data at VILLAGE from October, 2014, to January 2015. Valid Q data was recorded during "+"%.0f"%No_of_Storm_Intervals_DAM_Q+" events at FOREST, and "+"%.0f"%No_of_Storm_Intervals_LBJ_Q+" events at VILLAGE . Valid SSC data from T and Interpolated Grab samples was recorded during "+"%.0f"%No_of_Storm_Intervals_DAM_S+" events at FOREST, and "+"%.0f"%No_of_Storm_Intervals_LBJ_S+" events at VILLAGE. Of those storms, "+"%.0f"%len(S_Diff_table)+" events had valid P and SSY data for both the FOREST and VILLAGE to calculate and compare SSY from the UPPER and LOWER watersheds (Table "+S_Diff_table.table_num+"). Valid SSY data from Interpolated grab samples was collected at QUARRY for "+"%.0f"%No_of_Storm_Intervals_QUA_S+" storms to compare with SSY from FOREST and VILLAGE directly (Table "+S_Diff_table_quarry.table_num+"). Storm event durations ranged from "+"%.0f"%LBJ_StormIntervals['duration (hrs)'].min()+" hours to "+"%.0f"%max_storm_duration+" days, with mean duration of "+"%.0f"%LBJ_StormIntervals['duration (hrs)'].mean()+" hours.") 

#### Comparing SSY from disturbed and undisturbed subwatersheds
document.add_heading('Comparing SSY from disturbed and undisturbed subwatersheds',level=3)

Percent_Upper_S_2 = S_Diff_table['% UPPER']['Total/Avg:']
Percent_Lower_S = S_Diff_table['% LOWER']['Total/Avg:']
S_Diff_table_percents = S_Diff_table[S_Diff_table['% UPPER']!='-']
Percent_Upper_S_min, Percent_Upper_S_max =  "%0.1f"%S_Diff_table_percents['% UPPER'].astype(float).min(), "%.0f"%S_Diff_table_percents['% UPPER'].astype(float).max()
Percent_Lower_S_min, Percent_Lower_S_max =  "%.0f"%S_Diff_table_percents['% LOWER'].astype(float).min(), "%.0f"%S_Diff_table_percents['% LOWER'].astype(float).max()


document.add_paragraph("The Upper and Lower subwatersheds are similar in size, 0.90km2 and 0.88km2, so assuming the specific SSY is similar in both watersheds, they should account for roughly the same percentage sediment contribution to the total. However, for the "+"%.0f"%len(S_storm_diff_table(subset='post'))+" storms with valid data from both FOREST and VILLAGE, SSY from the UPPER subwatershed (SSYUPPER) ranged from "+Percent_Upper_S_min+"-"+Percent_Upper_S_max +"%, and accounted for an average of "+Percent_Upper_S_2+"% of Total SSY. SSY from the Lower watershed (SSYLOWER) ranged from "+Percent_Lower_S_min+"-"+Percent_Lower_S_max +"%, and accounted for an average of "+Percent_Lower_S+"% of Total SSY (Table "+S_Diff_table.table_num+"). This suggests that human disturbance in the Lower subwatershed has significantly increased Total SSY to Faga'alu Bay despite the relatively high amount of impervious surface associated with buildings and road surfaces in the village area.")

## Storm Sediment Table
if 'S_Diff_table' in locals():
    dataframe_to_table(df=S_Diff_table,table_num=S_Diff_table.table_num,caption="Sediment discharge from subwatersheds in Faga'alu",fontsize=9)

document.add_paragraph('')

Percent_UPPER_S_3 = S_Diff_table_quarry['% UPPER']['Total/Avg:']
Percent_QUARRY_S = S_Diff_table_quarry['% LOWER_QUARRY']['Total/Avg:']
Percent_VILLAGE_S = S_Diff_table_quarry['% LOWER_VILLAGE']['Total/Avg:']

SSY_UPPER_3, sSSY_UPPER_3 = S_Diff_table_quarry['UPPER tons']['Total/Avg:'], S_Diff_table_quarry['UPPER tons']['Tons/km2'] 
SSY_LOWER_QUARRY, sSSY_LOWER_QUARRY = S_Diff_table_quarry['LOWER_QUARRY tons']['Total/Avg:'], S_Diff_table_quarry['LOWER_QUARRY tons']['Tons/km2']
SSY_LOWER_VILLAGE, sSSY_LOWER_VILLAGE = S_Diff_table_quarry['LOWER_VILLAGE tons']['Total/Avg:'], S_Diff_table_quarry['LOWER_VILLAGE tons']['Tons/km2']
SSY_TOTAL_3, sSSY_TOTAL_3 = S_Diff_table_quarry['TOTAL tons']['Total/Avg:'], S_Diff_table_quarry['TOTAL tons']['Tons/km2']
P_measured_3, P_measured_3_perc_ann = S_Diff_table_quarry['Precip (mm)']['Total/Avg:'], "%.0f"%(float(S_Diff_table_quarry['Precip (mm)']['Total/Avg:'])/4000 *100)

Area_UPPER, Area_LOWER_QUARRY, Area_LOWER_VILLAGE = landcover_table.ix[0]['Area km2'], landcover_table.ix[1]['Area km2'], landcover_table.ix[2]['Area km2']
disturbed_area_LOWER_QUARRY = Area_LOWER_QUARRY * float(S_Diff_table_quarry['LOWER_QUARRY tons']['fraction disturbed (%)'])/100
disturbed_area_LOWER_VILLAGE = Area_LOWER_VILLAGE * float(S_Diff_table_quarry['LOWER_VILLAGE tons']['fraction disturbed (%)'])/100

document.add_paragraph("SSYEV data measured at FG2 was available for "+"%.0f"%No_of_Storm_Intervals_QUA_S+" of the storms in Table "+S_Diff_table.table_num+", so SSYEV from the LOWER subwatershed including the quarry (SSYLOWER_QUARRY) and the village areas below the quarry (SSYLOWER_VILLAGE) could be calculated to determine the relative sediment contribution from these sources (Table "+S_Diff_table_quarry.table_num+").")

if 'S_Diff_table_quarry' in locals():
    dataframe_to_table(df=S_Diff_table_quarry,table_num=S_Diff_table_quarry.table_num,caption="Sediment discharge from subwatersheds in Faga'alu",fontsize=9)
#document.add_paragraph("Storm on 2/3/12 has a potential outlier at DT at beginning of storm, which makes the SSYquarry huge! Storm at 2/5/12 doesn't have adequate SSC samples for quarry and its a multipeaked event so the SSY doesn't fall back to low levels like the LBJ and DAM T data suggest it should. Storm 3/6/13 looks like may have missed the second peak but the data looks comparable between sites. Storm 4/16/13 doesn't have alot of points but maybe they're adequate? Storm 4/23/13 looks good. Storm 4/30/13 has inadequate SSC data for all locations after the first peak; can maybe change the storm interval? Storm 6/5/13 has inadequate SSC data for all locations for the first peak, decent data for LBJ and DT for second peak but not for DAM; can maybe change the storm interval? Storm 2/14/14 looks good. Storm 2/20/14 looks good. Storm 2/21/14 looks good. Storm 2/27/14 looks kinda shitty. So good storms are: 3/6/13, 4/16/13? 4/23/13, 4/30/13?, 6/5/13?, 2/14/14, 2/20/14, 2/21/14") 

document.add_paragraph("For the measured storms in Table "+S_Diff_table_quarry.table_num+", TOTAL SSY was "+SSY_TOTAL_3+" tons, comprised of an average of "+Percent_UPPER_S_3+"% from the UPPER subwatershed, "+Percent_QUARRY_S+"% from LOWER_QUARRY subwatershed, and "+Percent_VILLAGE_S+"% from the LOWER_VILLAGE subwatershed. sSSY from the UPPER, LOWER_QUARRY, and LOWER_VILLAGE subwatersheds and TOTAL watershed was "+sSSY_UPPER_3+", "+sSSY_LOWER_QUARRY+", "+sSSY_LOWER_VILLAGE+", and "+sSSY_TOTAL_3+", respectively. sSSY from LOWER_QUARRY and LOWER_VILLAGE was "+S_Diff_table_quarry['LOWER_QUARRY tons']['DR']+" and "+S_Diff_table_quarry['LOWER_VILLAGE tons']['DR']+" times higher, respectively, than sSSY from UPPER subwatershed, suggesting human disturbance has significantly increased SSY over natural levels, particularly at the quarry. sSSY from the TOTAL watershed was "+S_Diff_table_quarry['TOTAL tons']['DR']+" times higher than the UPPER subwatershed, similar to the larger range of storms in Table "+S_Diff_table.table_num+", where specific SSY was "+S_Diff_table['TOTAL tons']['DR']+" times higher.")



  
    
#### Fitting SSY models
document.add_heading('Predicting SSYEV from storm metrics',level=3)


if 'SSY_models_ALL' in locals():
    document.add_picture(SSY_models_ALL['filename']+'.png',width=Inches(6))
    add_figure_caption(SSY_models_ALL['fig_num'],"SSY rating curves for predictors")
## Power law models from ALLStorms_ALLRatings 
PS_upper,PS_total,EI_upper,EI_total, QsumS_upper,QsumS_total,QmaxS_upper,QmaxS_total=ALLStorms_ALLRatings

# Pearson's and Spearman's correlation coeffs and r2 and RMSE
#document.add_paragraph("Pearson's correlation coefficients...")
if 'SSYEV_models_stats' in locals():
    dataframe_to_table(df=SSYEV_models_stats,table_num=SSYEV_models_stats.table_num,caption="Model statistics")

## Assessing alpha and Beta model parameters
ratings_table = ALLRatings_table(subset='post')
alpha_upper, beta_upper = ratings_table['alpha'].loc['Qmax_upper'], ratings_table['Beta'].loc['Qmax_upper']
alpha_total, beta_total = ratings_table['alpha'].loc['Qmax_total'], ratings_table['Beta'].loc['Qmax_total']

#### CONCLUSION
conclusion_title=document.add_heading('Conclusion',level=2)



#### Appendix
document.add_page_break()
document.add_heading('APPENDIX',level=2)
## Storm Water Discharge
if 'Q_Diff_table' in locals():
    dataframe_to_table(df=Q_Diff_table,table_num=Q_Diff_table.table_num,caption="Water discharge from subwatersheds in Faga'alu",fontsize=9)
## 
   
## Save Document
document.save(maindir+'Manuscript/ChangeOverTime/DRAFT-Fagaalu_Sediment_Yield_Post-mitigation.docx')

## Clean up any open figures
plt.close('all')








