# -*- coding: utf-8 -*-
"""
Created on Tue Feb 17 10:50:00 2015

@author: Alex
"""
#plt.ioff()
plt.close('all')
from docx import *
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
##  Create Document
document = Document(maindir+'Manuscript/JOH-template.docx')

#### Formatting Styles
## Body Style
document.styles['Normal'].font.name = 'Times'
document.styles['Normal'].font.size = Pt(12)
document.styles['Normal'].paragraph_format.first_line_indent = Inches(0.5)
document.styles['Normal'].paragraph_format.line_spacing = 1

#tables.styles['Normal'].font.name = 'Times'
#tables.styles['Normal'].font.size = Pt(12)
#tables.styles['Normal'].paragraph_format.first_line_indent = Inches(0.0)

# Heading 1
document.styles['Heading 1'].paragraph_format.first_line_indent = Inches(0.0)
document.styles['Heading 1'].font.name = 'Times'
document.styles['Heading 1'].font.size = Pt(14)
document.styles['Heading 1'].font.bold = True
# Heading 2
document.styles['Heading 2'].paragraph_format.first_line_indent = Inches(0.0)
document.styles['Heading 2'].font.name = 'Times'
document.styles['Heading 2'].font.size = Pt(13)
document.styles['Heading 2'].font.bold = True
# Heading 3
document.styles['Heading 3'].paragraph_format.first_line_indent = Inches(0.0)
document.styles['Heading 3'].font.name = 'Times'
document.styles['Heading 3'].font.size = Pt(12)
document.styles['Heading 3'].font.bold = True
# Heading 4
document.styles['Heading 4'].paragraph_format.first_line_indent = Inches(0.0)
document.styles['Heading 4'].font.name = 'Times'
document.styles['Heading 4'].font.size = Pt(12)
document.styles['Heading 4'].font.bold = True
# Heading 5
document.styles['Heading 5'].paragraph_format.first_line_indent = Inches(0.0)
document.styles['Heading 5'].font.name = 'Times'
document.styles['Heading 5'].font.size = Pt(12)
document.styles['Heading 5'].font.bold = True
document.styles['Heading 5'].font.italic = True

# Captions
document.styles['Caption'].paragraph_format.first_line_indent = Inches(0.0)

######## SOME TOOLS
figure_captions = []
def add_figure_caption(fig_num=str(len(document.inline_shapes)),caption=''):
    #manuscript_cap = document.add_paragraph("Insert Figure "+fig_num+" here")
    manuscript_cap = document.add_paragraph("Figure "+fig_num+". "+caption)
    manuscript_cap.paragraph_style = 'caption'
    manuscript_cap.paragraph_format.first_line_indent = Inches(0.0)
    #figs_cap = tables_and_figures.add_paragraph("Figure "+fig_num+". "+caption)
    figure_captions.append("Figure "+fig_num+". "+caption)
    return

table_titles = []
def dataframe_to_table(df=pd.DataFrame(),table_num=str(len(document.tables)+1),caption='',fontsize=11):
    ## Add a place holder in the manuscript
    manuscript_cap = document.add_paragraph("Insert Table "+table_num+" here")
    manuscript_cap = document.add_paragraph("Table "+table_num+". "+caption)
    manuscript_cap.paragraph_style = 'caption'
    manuscript_cap.paragraph_format.first_line_indent = Inches(0.0)
    
    ## construct table
    try:
        table = tables.add_table(rows=1, cols=len(df.columns)) 
        ## Merge all cells in top row and add caption text
        table_caption = table.rows[0].cells[0].merge(table.rows[0].cells[len(df.columns)-1])
        table_caption.text = "Table "+table_num+". "+caption
        ## Add  header
        header_row = table.add_row().cells
        col_count =0 ## counter  to iterate over columns
        for col in  header_row:
            col.text = df.columns[col_count] #df.columns[0] is the index
            col_count+=1
        ## Add data by  iterating over the DataFrame rows, then using a dictionary of DataFrame column labels to extract data
        col_labels = dict(zip(range(len(df.columns)),df.columns.tolist())) ## create dictionary where '1  to  n' is key for DataFrame columns
        for row in df.iterrows():  ## iterate over  rows in  DataFrame
            #print row[1]
            row_cells = table.add_row().cells ## Add a row to the  table
            col_count  = 0
            for cell in row_cells: ## iterate over the columns in the row
                #print row[1][str(col_labels[col_count])]
                cell.text = str(row[1][str(col_labels[col_count])]) ## and plug in data using a dictionary to  get column labels for DataFrame
                col_count+=1
        ## Format Table Style  
        table.style = 'TableGrid' 
        table.style.font.size = Pt(fontsize)
        table.autofit
        #table.num = str(len(document.tables)+1)
    except:
        pass
    table_titles.append("Table "+table_num+". "+caption)
    tables.add_paragraph("")
    return table
    
def add_equation(eq_table):
    t = document.add_table(rows=len(eq_table.rows),cols=len(eq_table.columns))
    t.rows[0].cells[1].text = eq_table.rows[0].cells[1].text
    t.rows[0].cells[2].text = eq_table.rows[0].cells[2].text   
    if len(eq_table.rows)>1:
        t.rows[1].cells[0].merge(t.rows[1].cells[2]).text = eq_table.rows[1].cells[0].text
    t.style = 'TableGrid'   
    t.autofit
    return t
    
###################################################################################################################################################################    


#### TABLES ########################################################################################################################################################
table_count=0
def tab_count():
    global table_count
    table_count+=1
    return str(table_count)
### Landcover_Table
landcover_table.table_num = str(tab_count())

### SSYEV budget for storms with FG1 and FG3 data
S_budget_table.table_num = str(tab_count())
S_budget_analysis_table.table_num = str(tab_count())

### SSYEV budget for storms with FG1, FG2, and FG3 data
S_budget_2_table.table_num = str(tab_count())
S_budget_2_analysis_table.table_num = str(tab_count())

### Model statistics table
All_Models_stats_table.table_num = str(tab_count())

### Annual SSY
est_Annual_SSY_table.table_num = str(tab_count())

### SSY literature comparision
lit_values_for_Annual_ssy_table_num = str(tab_count())

#### FIGURES ########################################################################################################################################################
figure_count=0
def fig_count():
    global figure_count
    figure_count+=1
    return str(figure_count)
## INTRODUCTION
#### Study Area Map
Study_Area_map = {'filename':maindir+'Figures/Maps/FagaaluInstruments land only map with regional.tif', 'fig_num':str(fig_count())}
#### Quarry_picture
Quarry_picture = {'filename':maindir+'Figures/Quarry_Mitgation/Quarry before retention ponds.tif','fig_num':str(fig_count())}

## METHODS
#### Stage-Discharge Rating Curves
## LBJ_StageDischarge
LBJ_StageDischarge = {'filename':figdir+"Q/Water Discharge Ratings for FG3 (LBJ)",'fig_num':str(fig_count())}
plotQratingLBJ(ms=8,show=False,log=False,save=True,filename=LBJ_StageDischarge['filename'])
## DAM_StageDischarge
DAM_StageDischarge = {'filename':figdir+"Q/Water Discharge Ratings for FG1 (DAM)",'fig_num':str(fig_count())}
plotQratingDAM(ms=8,show=False,log=False,save=True,filename=DAM_StageDischarge['filename'])

#### T-SSC Rating Curves
T_SSC_Rating_Curves = {'filename':figdir+'T/T-SSC rating curves','fig_num':str(fig_count())}
plot_all_T_SSC_ratings(Use_All_SSC=False,storm_samples_only=True,show=False,save=True,filename=T_SSC_Rating_Curves['filename'])

## RESULTS
#### Q time series
Q_timeseries =  {'filename':figdir+"Q/Q time series 2012-2014",'fig_num':str(fig_count())}
plot_Q_by_year(log=False,show=False,save=True,filename=Q_timeseries['filename'])

#### Individual Storm
Example_Storm = {'filename':figdir+'storm_figures/Example_Storm','fig_num':str(fig_count())}
example_storm_start, example_storm_end = All_Storms.ix[143]['start'], All_Storms.ix[143]['end']
plot_storm_individually(storm_data[example_storm_start:example_storm_end],show=False,save=True,filename=Example_Storm['filename']) 


####  SSC Boxplots
SSC_Boxplots= {'filename':figdir+'SSC/Grab sample boxplots baseflow and stormflow','fig_num':str(fig_count())}
## returns f,p and t-test statistics
f1,p1,QUARRY_DAM_ttest1,QUARRY_LBJ_ttest1,H1,KWp1,QUARRY_DAM_mannwhit1,QUARRY_LBJ_mannwhit1, f2,p2,QUARRY_DAM_ttest2,QUARRY_LBJ_ttest2,H2,KWp2,QUARRY_DAM_mannwhit2,QUARRY_LBJ_mannwhit2 = SSC_box_plots(subset=['Pre-baseflow','Pre-storm'],withR2=False,log=True,show=False,save=True,filename=SSC_Boxplots['filename']) 

#### Q vs. SSC
Discharge_Concentration = {'filename':figdir+'SSC/Water discharge vs Sediment concentration','fig_num':str(fig_count())}
plotQvsC(subset=['Pre-baseflow','Pre-storm'],ms=6,show=False,log=True,save=True,filename=Discharge_Concentration['filename'])

#### SSY models
SSY_models_ALL = {'filename':figdir+'SSY/SSY Models ALL pre-mitigation','fig_num':str(fig_count())}
All_Storms_All_Models = plot_All_Storms_All_Models(subset='pre',ms=4,norm=True,log=True,show=True,save=True,filename=SSY_models_ALL['filename'])

###### EQUATIONS ############################################################################################################################################
equation_count=0
def eq_count():
    global equation_count
    equation_count+=1
    return str(equation_count)
    
Equations = Document(maindir+'/Manuscript/Equations-9_17_15.docx').tables

## SSYev = Q*SSC 
SSYEV_eq = Equations[0].table
SSYEV_eq.eq_num = eq_count()

## SSY disturbed
SSY_disturbed = Equations[1].table
SSY_disturbed.eq_num = eq_count()

## DR = SSY/SSYPRE
DR_eq = Equations[2].table
DR_eq.eq_num = eq_count()

## predict_SSYev = aXb
predict_SSYEV_eq = Equations[3].table
predict_SSYEV_eq.eq_num = eq_count()

## SSYannual
SSY_annual_eq = Equations[4].table
SSY_annual_eq.eq_num = eq_count()

## PE = sqrt(sum(Error^2+Error^2))
PE_eq = Equations[5].table
PE_eq.eq_num = eq_count()

############################################################################################################################################
#### Appendix
table_count,figure_count,equation_count=0, 0, 0
### Storm Water Discharge Table
Q_budget_table.table_num =str(tab_count())

## Cross-Sections
#LBJ
LBJ_Cross_Section = {'filename':figdir+'Q/LBJ_Cross_Section','fig_num':'A1.'+str(fig_count())}
Mannings_Q_from_CrossSection(datadir+'Q/Cross_Section_Surveys/LBJ_cross_section.xlsx','LBJ_m',Slope=0.016,Manning_n='Jarrett',k=.06/.08,stage_start=1.4,show=False,save=True,filename=LBJ_Cross_Section['filename'])

#DAM
DAM_Cross_Section = {'filename':figdir+'Q/DAM_Cross_Section','fig_num':'A1.'+str(fig_count())}
Mannings_Q_from_CrossSection(datadir+'Q/Cross_Section_Surveys/DAM_cross_section.xlsx','DAM_m',Slope=0.016,Manning_n='Jarrett',k=.04/.07,stage_start=.46,show=False,save=True,filename=DAM_Cross_Section['filename'])

## Synthetic Rating Curves
Synthetic_Rating_Curve = {'filename':figdir+'T/Synthetic Rating Curves Fagaalu','fig_num':'A3.'+str(fig_count())} ## define file name to find the png file from other script
#Synthetic_Rating_Curves_Fagaalu(param='SS_Mean',show=False,save=True,filename=Synthetic_Rating_Curve['filename'])## generate figure from separate script and save to file


############################################################################################################################################
#### TITLE
title_title = document.add_heading('TITLE:',level=1)
title_title.paragraph_format.space_before = 0
title = document.add_heading('Contributions of human activities to suspended sediment yield during storm events from a small, steep, tropical watershed',level=1)
title.paragraph_format.space_before = 0

## subscript/superscript words
# Sub
p = document.add_paragraph("SSY"); p.add_run("EV").font.subscript = True
p.add_run(" SSY"); p.add_run("UPPER").font.subscript = True
p.add_run(" sSSY"); p.add_run("UPPER").font.subscript = True
p.add_run(" SSY"); p.add_run("LOWER").font.subscript = True
p.add_run(" sSSY"); p.add_run("LOWER").font.subscript = True
p.add_run(" SSY"); p.add_run("TOTAL").font.subscript = True
p.add_run(" sSSY"); p.add_run("UPPER").font.subscript = True
p.add_run(" SSY"); p.add_run("FG3").font.subscript = True
p.add_run(" SSY"); p.add_run("FG2").font.subscript = True
p.add_run(" SSY"); p.add_run("FG1").font.subscript = True
p.add_run(" SSY"); p.add_run("LOWER_QUARRY").font.subscript = True
p.add_run(" SSY"); p.add_run("LOWER_VILLAGE").font.subscript = True
p.add_run(" SSY"); p.add_run("disturbed").font.subscript = True
p.add_run(" SSY"); p.add_run("subwatershed").font.subscript = True
p.add_run(" Area"); p.add_run("undist").font.subscript = True
p.add_run(" E"); p.add_run("Qmeas").font.subscript = True
p.add_run(" E"); p.add_run("SSCmeas").font.subscript = True

## AUTHORS
document.add_heading('Authors:',level=3)
p = document.add_paragraph("Messina, A.M.")
p.add_run("a*").font.superscript = True
p.add_run(", Biggs, T.W.")
p.add_run("a").font.superscript = True
p.paragraph_format.first_line_indent = Inches(0.0)

p = document.add_paragraph("")
p.add_run("a").font.superscript = True
p.add_run(" San Diego State University, Department of Geography, San Diego, CA 92182, amessina@rohan.sdsu.edu, +1-619-594-5437, tbiggs@mail.sdsu.edu, +1-619-594-0902")
p.paragraph_format.first_line_indent = Inches(0.0)
document.add_paragraph("")

## Annual SSY
min_annual_SSY_UPPER, max_annual_SSY_UPPER = est_Annual.loc['UPPER'].min(), est_Annual.loc['UPPER'].max()
min_annual_SSY_TOTAL, max_annual_SSY_TOTAL = est_Annual.loc['TOTAL'].min(), est_Annual.loc['TOTAL'].max()
## Annual sSSY
min_annual_sSSY_UPPER, max_annual_sSSY_UPPER = est_Annual.loc['UPPER.'].min(), est_Annual.loc['UPPER.'].max()
min_annual_sSSY_TOTAL, max_annual_sSSY_TOTAL = est_Annual.loc['TOTAL.'].min(), est_Annual.loc['TOTAL.'].max()

#### ABSTRACT
abstract_title = document.add_heading('ABSTRACT',level=2)
abstract_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract = document.add_paragraph("Anthropogenic watershed disturbance by deforestation, mining, agriculture, and urbanization often increases fluvial sediment yields, enhancing sediment stress on aquatic ecosystems near the outlets of impacted watersheds. Suspended sediment yields (SSY) from undisturbed and human-disturbed portions of a small (1.8 km"u"\u00B2""), mountainous watershed that drains to a sediment-stressed coral reef were measured during storm and non-storm periods. Event-wise SSY (SSYEV) for "+str(len(Storms_LBJ['Qsum'].dropna()))+" storms was calculated from measurements of water discharge (Q), turbidity (T), and suspended sediment concentration (SSC) collected downstream of three key sediment sources: undisturbed forest, an aggregate quarry, and an urban area. SSC and SSYEV were significantly higher downstream of the quarry during both storm- and non-storm periods. The human-disturbed subwatershed accounted for an average of "+S_budget_2.loc['Total/Avg','% LOWER']+"-"+S_budget.loc['Total/Avg','% LOWER']+"% of SSYEV from the total watershed, and has increased loads to the coast by "+S_budget_2.loc['DR','TOTAL tons']+"-"+S_budget.loc['DR','TOTAL tons']+"x over natural background. Specific SSY (tons/area) from the disturbed quarry area was "+S_budget_2_analysis.loc['DR for sSSY from disturbed areas','LOWER_QUARRY']+"x higher than natural forest, and the quarry, which covers 1.1% of the total watershed area, contributed "+S_budget_2.loc['Total/Avg','% LOWER_QUARRY']+"% of total SSYEV at the outlet. Similar to mountainous watersheds in semi-arid and temperate climates, SSYEV from both the undisturbed and disturbed watersheds showed strong correlation with maximum event discharge (Qmax, Pearson's r = "+All_Models_stats.loc['Qmax_upper']['Pearson']+" and "+All_Models_stats.loc['Qmax_total']['Pearson']+" for the undisturbed and disturbed watersheds) as well as event total precipitation, and event total Q, but not with an erosivity index. Annual sediment yield estimates varied from "+min_annual_SSY_UPPER+"-"+max_annual_SSY_UPPER+" tons/yr ("+min_annual_sSSY_UPPER+"-"+max_annual_sSSY_UPPER+" tons/km"u"\u00B2""/yr) from the undisturbed subwatershed, and "+min_annual_SSY_TOTAL+"-"+max_annual_SSY_TOTAL+" tons/yr ("+min_annual_sSSY_TOTAL+"-"+max_annual_sSSY_TOTAL+" tons/km"u"\u00B2""/yr) from the human-disturbed subwatershed, depending on the estimation method. Only 5.2% of the watershed is disturbed by humans but sediment yield has been increased significantly ("+S_budget.loc['DR','TOTAL tons']+"x). Sediment loads were very sensitive to land cover change, and identification of sediment hotspots like the quarry will help sediment mitigation and coral restoration efforts.")

#### KEYWORDS
document.add_heading('Keywords:',level=2)
keywords = document.add_paragraph("Sediment yield, volcanic islands, mountainous catchments, land use, storm events, coastal sediment load, American Samoa")
keywords.paragraph_format.first_line_indent = Inches(0.0)

terms = "     =  3   )    "

#### INTRODUCTION
introduction_title = document.add_heading('Introduction',level=2)

## What's the problem? Where is the problem? How is it being addressed?
document.add_paragraph("Human activities including deforestation, agriculture, road construction, mining, and urbanization alter the timing, composition, and amount of sediment loads to downstream ecosystems (Syvitski et al., 2005). Increased sediment loads can stress aquatic ecosystems, including coral reefs that occur near the outlets of impacted watersheds. Sediment impacts coral by decreasing light for photosynthesis and increasing sediment accumulation rates (Fabricius, 2005; Storlazzi et al., 2015; West and van Woesik, 2001). Anthropogenic sediment disturbance can be particularly high on volcanic islands in the humid tropics, which have a high potential for erosion due to high rainfall, extreme weather events, steep slopes, and erodible soils. Sediment yield in densely vegetated watersheds can be particularly sensitive to land clearing, which alters the fraction of exposed soil more than in sparsely-vegetated regions. The steep topography and small floodplains in small volcanic islands further limits sediment storage and the capacity of the watershed to buffer increased sediment yields. Such environments characterize many volcanic islands in the south Pacific, which also contain many coral reefs impacted by sediment.")

document.add_paragraph("A large proportion of a watershed's sediment yield can originate in hotspots, which are disturbed areas that cover a relatively small fraction of the watershed area. In the Caribbean, unpaved roads covering 0.3-0.9% of the watershed area were the dominant sediment source in disturbed watersheds on St. John, and increased sediment yield to the coast by 5-9 times, relative to undisturbed watersheds (Ramos-Scharr"u"\u00F3""n and Macdonald, 2007). In the Pacific Northwest of the United States, several studies found most road-generated sediment can originate from just a small fraction of unpaved roads (Gomi et al., 2005; Henderson and Toews, 2001; Megahan et al., 2001; Wemple et al., 1996), and heavily used roads could generate 130 times as much sediment as abandoned roads (Reid and Dunne, 1984). In a watershed disturbed by grazing on Molokai, Hawaii, less than 5% of the land produces most of the sediment, and only 1% produces approximately 50% of the sediment (Risk, 2014; Stock et al., 2010), suggesting that management should focus on identifying, quantifying, and mediating erosion hotspots.") 

document.add_paragraph("Management of sediment requires linking land use changes and mitigation strategies to changes in sediment yields at the watershed outlet. A sediment budget quantifies sediment as it moves from key sources to its eventual exit from a watershed (Rapp, 1960), and is useful to characterize watershed response to land use change and management interventions (Walling and Collins, 2008). Walling (1999) used a sediment budget to show that sediment yield from watersheds can be insensitive to both land use change and erosion management due to high sediment storage capacity on hillslopes and in the channel. Sediment yield from disturbed areas can be large but may not be important compared to naturally high yields from undisturbed areas. While a full description of all sediment production and transport processes are of scientific interest, the sediment budget needs to be simplified to be used as a management tool (Slaymaker, 2003). Most management applications require only that the order of magnitude or the relative importance of process rates be known, so Reid and Dunne (1996) argue a management-focused sediment budget can be developed quickly in situations where the management problem is clearly defined and the management area can be divided into homogenous sub-units.")

document.add_paragraph("Knowledge of suspended sediment yield (SSY) under both natural and disturbed conditions on most tropical, volcanic islands remains limited, due to the challenges of in situ monitoring. Existing sediment yield models are often designed for agricultural landscapes and are not well-calibrated to the climatic, topographic, and geologic conditions found on steep, tropical islands. Most readily available models also do not incorporate many of the important processes that generate sediment in steep watersheds, including mass movements (Calhoun and Fletcher, 1999; Ramos-Scharr"u"\u00F3""n and Macdonald, 2005; Sadeghi et al., 2007). Developing models that predict SSY from small, mountainous catchments is a significant contribution for establishing baselines for change-detection for sediment mitigation projects, and can also further improve models applied at the regional scale (Duvert et al., 2012).") 

document.add_paragraph("Traditional approaches to quantifying human impact on sediment budgets, including comparison of total annual yields (Fahey et al., 2003) and sediment rating curves (Asselman, 2000; Walling, 1977), are complicated by interannual variability and hysteresis in the discharge-concentration relationship. As an alternative, other studies have compared SSY generated by storm events of the same magnitude to assess the contribution of individual subwatersheds to total SSY (Zimmermann et al., 2012), compare the responses of different watersheds to "+'"storm metrics"'+" (Basher et al., 2011; Duvert et al., 2012; Fahey et al., 2003; Hicks, 1990), and determine changes in SSY from the same watershed over time (Bonta, 2000).")

document.add_paragraph("Event-wise SSY (SSYEV) may correlate with various precipitation or discharge variables ("+'"storm metrics"'+"), such as total precipitation, the Erosivity Index (Kinnell, 2013), or total discharge, but the best correlation has consistently been found with maximum event discharge (Qmax). Several researchers have hypothesized that Qmax integrates the hydrological response of a watershed, making it a good predictor of SSYEV in diverse environments (Duvert et al., 2012; Rankl, 2004). High correlation between SSYEV and Qmax has been found in semi-arid, temperate, and sub-humid watersheds in Wyoming (Rankl, 2004), Mexico, Italy, France (Duvert et al., 2012), and New Zealand (Basher et al., 2011; Hicks, 1990), but this approach has not been attempted for steep, tropical watersheds on volcanic islands.")

document.add_paragraph("The anthropogenic impact on SSYEV may vary by storm magnitude, as documented in Pacific Northwest forests (Lewis et al., 2001). As storm magnitude increases, water yield and/or SSYEV from natural areas may increase relative to human-disturbed areas, diminishing anthropogenic impact relative to the natural baseline. While large storms account for most SSY in natural conditions, human-disturbed areas may show the most significant disturbance for smaller storms (Lewis et al., 2001). The disturbance ratio (DR) may be highest for small storms, when background SSY from the undisturbed forest is low and erodible sediment from disturbed surfaces is the dominant source. For large storms, mass movements and bank erosion may contribute to naturally high SSY from the undisturbed watershed, reducing the DR for large events. Natural variations in climate, such as El Nino-Southern Oscillation events (Wulf et al. 2012) and the frequency of large-magnitude events (Wolman and Miller 1960) can also have large impacts on sediment discharge from undisturbed watersheds. Given the geologically short time scale of this study, it was assumed the sediment discharge from the undisturbed watershed is at a steady-state. ") 

document.add_paragraph("This study uses in situ measurements of precipitation (P), stream discharge (Q), turbidity (T) and suspended sediment concentration (SSC) to 1) quantify suspended sediment yield from undisturbed and human-disturbed parts of a small watershed in the south Pacific and 2) to develop an empirical model of storm-generated suspended sediment yield. The questions addressed include: How much has human disturbance increased sediment yield to the coast?  What human activities dominate the anthropogenic contribution to the sediment budget? Which storm metric is the best predictor of storm event suspended sediment yield (SSYEV): total precipitation, Erosivity Index, total discharge, or maximum discharge? How do sediment contributions from human-disturbed areas and undisturbed areas vary with storm size?")

#### STUDY AREA
study_area_title = document.add_heading('Study Area',level=2)

document.add_paragraph("The study watershed, Faga'alu, is located on Tutuila (14S, 170W), the largest island in the Territory of American Samoa (140 km"u"\u00B2""). Like many volcanic islands in the Pacific, Tutuila is composed of steep, heavily forested mountains with villages and roads mostly confined to the flat areas near the coast. The main stream in Faga'alu runs the length of the watershed (~3 km), and drains an area of 1.78 km"u"\u00B2"" (area draining to FG3 in Figure 1). The main watershed includes Matafao Mountain, the highest point on Tutuila (653 m), and the stream discharges into the Pacific Ocean. The mean slope of the main Faga'alu watershed is 0.53 m/m and total relief is 653 m. Several small ephemeral streams drain directly to the ocean (0.63 km"u"\u00B2"") (grey dotted boundary in Figure "+Study_Area_map['fig_num']+"). The stream discharges to an adjacent, fringing coral reef embayment that Fenner et al. (2008) identified as being highly degraded by sediment. Faga'alu watershed was identified by local environmental management agencies in the American Samoa Coral Reef Advisory Group (CRAG) as a heavily impacted watershed, and in August 2012 was selected by the US Coral Reef Task Force (USCRTF) as a Priority Watershed for conservation and remediation efforts (Holst-Rice et al., 2015).")

## Study Area map
if 'Study_Area_map' in locals():
    document.add_picture(Study_Area_map['filename'],width=Inches(6))
    add_figure_caption(Study_Area_map['fig_num'],"Faga'alu watershed showing the UPPER (undisturbed) and LOWER (human-disturbed) subwatersheds. The LOWER subwatershed drains areas between FG1 and FG3, and is further subdivided into the LOWER_QUARRY containing the quarry (between FG1 and FG2) and the LOWER_VILLAGE containing the village areas (between FG2 and FG3). The TOTAL watershed includes all subwatersheds draining to FG3. The Administrative watershed boundary for government jurisdiction is outlined by the dotted white line. Blue pentagons in the UPPER watershed show the location of abandoned water supply reservoirs (see Appendix 2 for full description). Barometer locations at NSTP6 and TULA shown in top-right.")
    
# Climate
document.add_heading('Climate',level=3)
document.add_paragraph("Precipitation on Tutuila is caused by several mechanisms including cyclones and tropical depressions, isolated thunderstorms, and orographic uplifting of trade-wind squalls over the high (300-600 m), mountainous ridge that runs the length of the island. Unlike many other Pacific Islands, the ridge runs parallel to the predominant wind direction, and does not cause a significant windward/leeward rainfall gradient. Average annual specific discharge (m"u"\u00B3""/yr/km"u"\u00B2"") shows little spatial variation across the island, irrespective of watershed location or orientation (Dames & Moore, 1981). From 1903 to 1973, average annual precipitation over the island was 3,800 mm/yr (Eyre, 1989; Izuka et al., 2005). Precipitation increases with elevation, from an average 2,380 mm/yr at the shoreline to 6,350 mm/yr at the highest elevation on the island. In Faga'alu watershed, rainfall records show average annual precipitation is 6,350 mm at Matafao Mtn. (653 m m.a.s.l), 5,280 mm at Matafao Reservoir (249 m m.a.s.l.) and about 3,800 mm on the coastal plain (Craig, 2009; Dames & Moore, 1981; Perreault, 2010; Tonkin & Taylor International Ltd., 1989; Wong, 1996). Mean annual potential evapotranspiration follows the opposite trend, varying from 890 mm at high elevation to 1,150 mm at sea level (Izuka et al., 2005). Tropical cyclones are erratic but occurred on average every 1-13 years from 1981-2014 (Craig, 2009) and bring intense rainfall, flooding, landslides, and high sediment yield (Buchanan-Banks, 1979).")

document.add_paragraph("There are two subtle rainfall seasons: a drier winter season, from June through September and a wetter summer season, from October through May (Izuka et al., 2005). During the drier winter season, the island is influenced by relatively stronger, predominantly east to southeast tradewinds, lower temperatures, lower humidity and lower total rainfall. During the wetter summer season the Inter-Tropical Convergence Zone (ITCZ) moves over the region, causing light to moderate Northerly winds, higher temperatures, higher humidity, and higher total rainfall. While total rainfall is lower in the drier tradewind season, large storm events are still observed. Analysis of 212 peak discharges at 11 continuous-record gaging sites 1959-1990 showed 65% of annual peak flows occurred during the wet season and 35% of annual peak flows occurred during the drier Tradewind season (Wong, 1996). Analysis of mean monthly rainfall data for the period 1971-2000 showed that 75% of precipitation occurred in the wet season, which includes 67% of the year (October-May), and 25% occurred in the dry season, which covers 33% of the year (June-September) (Perreault, 2010; data from USGS rain gauges and Parameter-elevation Relationships on Independent Slopes Model (PRISM) Climate Group (Daly et al., 2008)).")

# Land Use
document.add_heading('Land Cover and Land Use',level=3)
document.add_paragraph("The predominant land cover in Faga'alu watershed is undisturbed vegetation (94.8%), including forest (85.7%) and scrub/shrub (9.0%) on the steep hillsides (Table 1), based on a land cover map from NOAA's Ocean Service and Coastal Services Center (2010). The upper watershed is dominated by undisturbed rainforest on steep hillslopes. The lower subwatershed has steep vegetated hillslopes and a relatively small flat area in the valley bottom that is urbanized. This settlement pattern is typical in the South Pacific and other volcanic islands, where their small size and steep topography constrain development to valley bottoms near the coast (B"u"\u00E9""gin et al., 2014).  Compared to other watersheds on Tutuila, a relatively large portion of Faga'alu watershed is urbanized (3.2% "+'"High Intensity Developed"'+" in Table 1), due to large areas of impervious surface associated with the hospital and the numerous residences and businesses. A small portion of the watershed (0.9%) is developed open space, which includes landscaped lawns and parks. In addition to some small household gardens there are several small agricultural areas of banana and taro on the steep hillsides. These agricultural plots were classified as grassland due to the high fractional grass cover in the plots. Farmers of these plots receive technical assistance from the Natural Resource Conservation Service (NRCS) to mitigate erosion. There are several small footpaths and unpaved driveways in the village, but most unpaved roads are stabilized with compacted gravel and do not appear to be a major contributor of sediment (Horsley-Witten, 2012). Longitudinal sampling of Faga'alu stream during low flow conditions in 2011 showed significantly increased turbidity downstream of a bridge construction site on the village road approximately 200 m downstream of FG2 (Curtis et al., 2011). Construction of the bridge was completed in March 2012 and no longer increases turbidity.")

    
## Quarry description    
document.add_paragraph("An open-pit aggregate quarry covers 1.6 ha and accounts for the majority of the bare land, which covers 1.1% of the Faga'alu watershed (Table 1). The quarry has been in continuous operation since the 1960's by advancing into the steep hillside to quarry the underlying basalt formation (Latinis et al., 1996). The overburden of soil and weathered rock was either piled up on-site where it was eroded by storms, or was manually rinsed from crushed aggregate. With few sediment runoff controls in place, sediment was discharged directly to Faga'alu stream. In 2011, the quarry operators installed some sediment runoff management practices such as silt fences and settling ponds (Horsley-Witten, 2011) but they were unmaintained and inadequate to control the large amount of sediment mobilized during storm events (Horsley-Witten, 2012). During the study period (2012-2014), additional sediment control measures were installed and some large piles of overburden were naturally overgrown by vegetation (Figure 2), altering the sediment availability. In late 2014, large sediment retention ponds were installed to mitigate sediment runoff and work is underway to document the reduction in sediment loading (Messina and Biggs, forthcoming; See Holst-Rice et al. (2015) for a full description of sediment mitigation efforts at the quarry).")

## Quarry picture
if 'Quarry_picture' in locals():
    document.add_picture(Quarry_picture['filename'],width=Inches(6))
    add_figure_caption(Quarry_picture['fig_num'],"Photos of the open aggregate quarry in Faga'alu in 2012, 2013, and 2014. Pictures a-b show vegetation overgrowth during the period of study from 2012-2014, and the location of the groundwater diversion that was installed in 2012. Pictures c-d show haul roads were covered in gravel in 2013 to limit fine sediment exposure. Photos: Messina")
    
document.add_paragraph("Three water impoundment structures were built in the early 20th century in the upper part of the watershed for drinking water supply and hydropower but only the highest, Matafao Reservoir, was ever connected to the municipal water system and has since fallen out of use (Tonkin & Taylor International Ltd., 1989)(Figure "+Study_Area_map['fig_num']+"). The dam at point FG1 has filled with bedload sediment and flows over the spillway even at the lowest flows. We assume the other reservoirs are similarly filled with coarse sediment and are not currently retaining fine suspended sediment. A full description of stream impoundments is in Appendix 2.")

#### METHODS
methods_title = document.add_heading('Methods',level=2)
document.add_paragraph("The suspended sediment yield (SSY) in Faga'alu stream was measured at three sampling points that drain key land covers we hypothesized would have different SSY: FG1 drains undisturbed forest in the UPPER subwatershed (watershed boundary to FG1), FG2 drains undisturbed forest and the quarry in the LOWER_QUARRY subwatershed (between FG1 and FG2), and FG3 drains undisturbed forest and the village in the LOWER_VILLAGE subwatershed (between FG2 and FG3) (Table 1). FG3 is also the watershed outlet for the TOTAL watershed.")

## Calculating SSYEV
document.add_heading('Calculating suspended sediment yield from individual storm events (SSYEV)',level=3)
document.add_paragraph("SSYEV at FG1, FG2, and FG3 were calculated by integrating continuous estimates of suspended sediment yield, calculated from measured or modeled water discharge (Q) and measured or modeled suspended sediment concentration (SSC) (Duvert et al., 2012):")

add_equation(SSYEV_eq) ## Equation 1

## Defining Storm Events


document.add_paragraph("Storm events can be defined by precipitation (Hicks, 1990) or discharge parameters (Duvert et al., 2012), and the method used to identify storm events on the hydrograph can significantly influence the analysis of SSYEV (Gellis, 2013). Dunne and Leopold (1978) assert that all hydrograph separation schemes are arbitrary and usually have little to do with the processes that generate storm flow, but if a consistent method is used then at least the results of different analyses can be compared. Graphical techniques may be implemented to separate the hydrograph into baseflow and quickflow, using the start and end of quickflow as the start and end of the storm event (Dunne and Leopold, 1978; Perreault, 2010). Storms can also be filtered from the analysis by using various criteria such as minimum storm duration, time between discharge peaks, minimum peak discharge, or more complex schemes using statistical distributions of flow percentiles (Lewis et al. 2001; Gellis 2013). More complex signal processing methods can also be used, including finding the inflection point of the second derivative of the hydrograph to determine the end of the storm event. However, complex events occur where subsequent precipitation generates stormflow before the stream has returned to baseflow, and the storm definition scheme can significantly affect the analysis of storm sediment yields by separating or combining multiple hydrograph peaks. Due to the high number of storm events and the prevalence of complex storm events at the study site, an automated approach that robustly separated complex events was desirable. The storm definition approach used in this study performed baseflow separation with a digital filter signal processing technique (Nathan and McMahon 1990) embedded in the R-statistical package EcoHydRology (Fuka et al. 2014). Only events with quickflow for at least one hour and peak flow greater than 10% of baseflow were included in the analysis. This approach was easily automated for application to a large number of events, and adequately separated complex storm events with multiple hydrograph peaks.")

### Relating sediment load to sediment budget
document.add_heading('Relationship of sediment load to sediment budget',level=3)
document.add_paragraph("We use the measured sediment yield at three locations to quantify the in-stream sediment budget. Other components of sediment budgets include channel erosion and or channel and floodplain deposition (Walling and Collins, 2008). Sediment storage and remobilization can significantly complicate the interpretation of in-stream loads, and complicate the identification of a land use signal. In Faga'alu, the channel bed is predominantly large volcanic cobbles and coarse gravel, with no significant deposits of fine sediment. Upstream of the village, the valley is very narrow with no floodplain. In the downstream reaches of the lower watershed, where fines might deposit in the floodplain, the channel has been stabilized with cobble reinforced by fencing, so overbank flows and sediment deposition on the floodplain are not observed. We therefore assume that channel erosion and channel and floodplain deposition are insignificant components of the sediment budget, so the measured sediment yields at the three locations reflect differences in hillslope sediment supply. Minimal sediment storage also reduces the lag time between landscape disturbance and observation of sediment at the watershed outlet.")

### Comparing SSY from disturbed and undisturbed subwatersheds
document.add_heading('Quantifying SSY from disturbed and undisturbed subwatersheds',level=3)
document.add_paragraph("A main objective for this study was to quantify anthropogenic changes in SSY from Faga'alu Stream (SSYTOTAL measured at FG3). Relative contributions to SSYTOTAL from undisturbed and human-disturbed areas were assessed using two approaches: 1) comparing SSY contributions from subwatersheds for each storm and the average of all storms, and 2) the Disturbance Ratio (DR).")

document.add_paragraph("The percent contributions of subwatersheds to SSYTOTAL were calculated from SSYEV measured at FG1, FG2, and FG3 (Figure 1). SSY from the UPPER subwatershed was measured at FG1 (SSYUPPER = SSYFG1). SSY from the LOWER subwatershed was calculated as SSYLOWER=SSYFG3-SSYFG1. Where SSYEV data at FG2 were also available, the contributions from the quarry subwatershed (SSYLOWER_QUARRY = SSYFG2-SSYFG1), and village subwatershed (SSYLOWER_VILLAGE = SSYFG3-SSYFG2) were calculated separately. ") 

document.add_paragraph("Land cover in the LOWER subwatershed includes both undisturbed and human-disturbed surfaces. To calculate SSY from disturbed areas, SSY from undisturbed areas was estimated using the specific SSY (sSSY tons/km"u"\u00B2"") from the UPPER subwatershed multiplied by the undisturbed area in the LOWER subwatershed:")

add_equation(SSY_disturbed) ## Equation

document.add_paragraph("The disturbance ratio (DR) is the ratio of SSYEV from the watershed under current condition  s to SSY under pre-disturbance conditions, estimated using sSSYUPPER:")

add_equation(DR_eq) ## Equation

document.add_paragraph("Both Equation "+SSY_disturbed.eq_num+" and "+DR_eq.eq_num+" assume that the whole watershed was originally covered in forest, and sSSY from forested areas in the LOWER subwatershed equals sSSY from the undisturbed UPPER watershed. SSY from the disturbed portions of the LOWER subwatershed (Equation "+SSY_disturbed.eq_num+") was used to calculate a DR for just the disturbed areas in the LOWER subwatershed.")

### Predicting event suspended sediment yield (SSYEV)
document.add_heading("Predicting event suspended sediment yield (SSYEV)",level=3)
document.add_paragraph("Four storm metrics were tested as predictors of SSYEV: total event precipitation (Psum), event Erosivity Index (EI30) (Hicks, 1990; Kinnell, 2013), total event water discharge (Qsum), and maximum event water discharge (Qmax) (Duvert et al., 2012; Rodrigues et al., 2013). SSYEV and the discharge metrics (Qsum and Qmax) were normalized by watershed area to compare different sized watersheds.")

document.add_paragraph("The relationship between SSYEV and storm metrics is often best fit by a power law function:")

add_equation(predict_SSYEV_eq) ## Equation

document.add_paragraph("The regression coefficients ("u"\u03B1"" and "u"\u03B2"") for the UPPER and TOTAL watersheds were tested for statistically significant differences using Analysis of Covariance (ANCOVA) (Lewis et al., 2001). A higher intercept ("u"\u03B1"") for the human-disturbed watershed indicates higher sediment yield for the same size storm event, compared to sediment yield from undisturbed areas. A difference in slope ("u"\u03B2"") would indicate the relative sediment contributions from the subwatersheds change with increasing storm size. If regression slopes for the UPPER and TOTAL watersheds are significantly different, it supports the conclusion that the effect of human-disturbance changes with storm size.")

### Annual estimates of SSY and sSSY
document.add_heading("Annual estimates of SSY and sSSY",level=3)

#### FOR QMAX SSYann, should it be all predicted or use model to fill in values (like in Hicks 1990,1994 (storm-rating method), Basher 1997, Nearing 2007) what about comparing different model estimates?
document.add_paragraph("Annual estimates of SSY and sSSY were used to compare watersheds with other literature. A continuous annual time-series of SSY was not possible at the study site due to the discontinuous field campaigns and failure of or damage to the turbidimeters during some months. A continuous record of water depth and Q was available for 2014, so the Qmax-SSYEV model (Eq 4) was used to predict SSY for all storms in 2014 (Basher et al., 1997). Sediment mitigation structures were installed at the quarry in October 2014, greatly reducing SSY from the LOWER_QUARRY subwatershed (unpublished data), so the Qmax-SSY relationship developed prior to the mitigation was used to calculate the annual pre-mitigation sediment yield. For storms with no Qmax data at FG3, Qmax was predicted from a linear regression between Qmax at FG1 and Qmax at FG3 for the study period (not shown).")

document.add_paragraph("Annual SSY and sSSY were also estimated by multiplying SSY from measured storms by the ratio of annual storm precipitation (Psann) to the precipitation measured during storms where SSY was measured (Psmeas):")

add_equation(SSY_annual_eq) ## Equation

document.add_paragraph("Most SSY is discharged during a few, relatively large events, and it is assumed that small events do not contribute significantly to annual SSY (Stock and Tribble, 2010). This method assumes that the sediment yield per mm of storm precipitation is constant over the year, and that the size distribution of storms has no effect on SSY, though there is some evidence that SSY increases exponentially with storm size (Lewis et al., 2001; Rankl, 2004).")


### Data Collection
document.add_heading('Data Collection',level=3)
document.add_paragraph("Data on precipitation (P), water discharge (Q), suspended sediment concentration (SSC) and turbidity (T) were collected during three field campaigns: January-March, 2012, February-July 2013, and January-March 2014, and several intervening periods of unattended monitoring  by instruments with data loggers. Field sampling campaigns were scheduled to coincide with the period of most frequent storms in the November-May wet season, though large storms were sampled throughout the year.")

### Precipitation
document.add_heading('Precipitation',level=4)
document.add_paragraph("Precipitation (P) was measured at three locations in Faga'alu watershed using Rainwise RAINEW tipping-bucket rain gages (RG1 and RG2) and a Vantage Pro Weather Station (Wx) (Figure "+Study_Area_map['fig_num']+"). Data at RG2 was only recorded January-March, 2012, to determine a relationship between elevation and precipitation in the LOWER subwatershed. The total event precipitation (Psum) and event Erosivity Index (EI30) were calculated using data from RG1, with data gaps filled by 15 min interval precipitation data from Wx.  While previous data suggest that precipitation increases with elevation, here we do not calculate watershed-mean precipitation, and instead use precipitation depth at RG1 to indicate the depth of rainfall during a storm event.  Most sheetwash and rill erosion, which depends on rainfall intensity and EI30, occurred at the quarry, near the location of RG1. Rainfall data from RG1 is therefore most representative of rainfall at the quarry.")

### Water Discharge
document.add_heading('Water Discharge',level=4) 
## Duvert 2010 used a PT at a bridge in Potrerillos
document.add_paragraph("Stream gaging sites were chosen to take advantage of an existing control structure (FG1) and a stabilized stream cross section (FG3)(Duvert et al, 2010). At FG1 and FG3, Q was calculated from 15 minute interval stream stage measurements, using a stage-Q rating curve calibrated to manual Q measurements made under baseflow and stormflow conditions (Figures "+LBJ_StageDischarge['fig_num']+" and "+DAM_StageDischarge['fig_num']+"). Stream stage was measured with non-vented pressure transducers (PT) (Solinst Levelogger or Onset HOBO Water Level Logger) installed in stilling wells at FG1 and FG3. Barometric pressure data collected at Wx were used to calculate stage from the pressure data recorded by the PT. Data gaps in barometric pressure from Wx were filled by data from stations at Pago Pago Harbor (NSTP6) and NOAA Climate Observatory at Tula (TULA) (Figure "+Study_Area_map['fig_num']+"). Priority was given to the station closest to the watershed with valid barometric pressure data. Barometric data were highly correlated and the data source made little (<1cm) difference in the resulting water level. Q was measured in the field by the area-velocity method (AV) using a Marsh-McBirney flowmeter to measure flow velocity and channel surveys measure cross-sectional area (Harrelson et al., 1994; Turnipseed and Sauer, 2010).")

document.add_paragraph("AV-Q measurements could not be made at high stages at FG1 and FG3 for safety reasons, so stage-Q relationships were constructed to estimate a continuous record of Q. At FG3, the channel is rectangular with stabilized rip-rap on the banks and bed (Appendix Figure A1.1). Recorded stage varied from 4 to 147 cm. AV-Q measurements (n= 14) were made from 30 to 1,558.0 L/sec, covering a range of stages from 6 to 39 cm. The highest recorded stage was much higher than the highest stage with measured Q so the rating could not be extrapolated by a power law. Stream conditions at FG3 fit the assumption for Manning's equation, so the stage-Q rating at FG3 was created using Manning's equation, calibrating Manning's n (0.067) to the Q measurements (Figure 3).")

## LBJ stage-discharge rating
if 'LBJ_StageDischarge' in locals():
   document.add_picture(LBJ_StageDischarge['filename']+'.png',width=Inches(6))
   add_figure_caption(LBJ_StageDischarge['fig_num'],"Stage-Discharge relationships for stream gaging site at FG3 for (a) the full range of observed stage and (b) the range of stages with AV measurements of Q. RMSE was "+"%.0f"%Manning_AV_rmse(LBJ_Man_reduced,LBJstageDischarge)[0]+" L/sec, or "+"%.0f"%Manning_AV_rmse(LBJ_Man_reduced,LBJstageDischarge)[2]+"% of observed Q.")
   
   
document.add_paragraph("At FG1, the flow control structure is a masonry ogee spillway crest of a defunct stream capture. The structure is a rectangular channel 43 cm deep that transitions abruptly to gently sloping banks, causing an abrupt change in the stage-Q relationship (Appendix Figure "+DAM_Cross_Section['fig_num']+"). At FG1, recorded stage height ranged from "+"%.0f"%DAM['stage(cm)'].min()+" to "+"%.0f"%DAM['stage(cm)'].max()+" cm, while area-velocity Q measurements (n= "+"%.0f"%len(DAMstageDischarge)+") covered stages from "+"%.0f"%DAMstageDischarge['stage(cm)'].min()+" to "+"%.0f"%DAMstageDischarge['stage(cm)'].max()+" cm. Since the highest recorded stage ("+"%.0f"%DAM['stage(cm)'].max()+" cm) was higher than the highest stage with measured Q ("+"%.0f"%DAMstageDischarge['stage(cm)'].max()+" cm), and there was a distinct change in channel geometry above 43 cm the rating could not be extrapolated by a power law. The flow structure did not meet the assumptions for using Manning's equation to predict flow so the HEC-RAS model was used (Brunner 2010). The surveyed geometry of the upstream channel and flow structure at FG1 were input to HEC-RAS, and the HEC-RAS model was calibrated to the Q measurements (Figure "+DAM_StageDischarge['fig_num']+"). While a power function fit Q measurements better than HEC-RAS for low flow, HEC-RAS fit better for Q above the storm threshold used in analyses of SSY (Figure "+DAM_StageDischarge['fig_num']+").")

## DAM stage-discharge rating
if 'DAM_StageDischarge' in locals():
    document.add_picture(DAM_StageDischarge['filename']+'.png',width=Inches(6))
    add_figure_caption(DAM_StageDischarge['fig_num'],"Stage-Discharge relationships for stream gaging site at FG1 for (a) the full range of observed stage and (b) the range of stages with AV measurements of Q. RMSE was "+"%.0f"%HEC_AV_rmse(DAM_HECstageDischarge,DAMstageDischarge)[0]+" L/sec, or "+"%.0f"%HEC_AV_rmse(DAM_HECstageDischarge,DAMstageDischarge)[2]+"% of observed Q. "+'"Channel Top"'+" refers to the point where the rectangular channel transitions to a sloped bank and cross-sectional area increases much more rapidly with stage. A power-law relationship is also displayed to illustrate the potential error that could result if inappropriate methods are used.")

document.add_paragraph("Water discharge at FG2 was calculated as the product of the specific water discharge from FG1 (m"u"\u00B3""/0.9 km"u"\u00B2"") and the watershed area draining to FG2 (1.17 km"u"\u00B2""). This assumes that specific water discharge from the subwatershed above FG2 is similar to above FG1. Discharge may be higher from the quarry surface, which represents "+"%.1f"%landcover_table.ix['LOWER_QUARRY (FG2)']['Bare (B)']+"% of the LOWER_QUARRY subwatershed, so Q, and thus SSY from the quarry are a conservative, lower bound estimate. The quarry surface is continually being disturbed, sometimes with large pits excavated and refilled in the course of weeks, as well as intentional water control structures implemented over time. Given the changes in the contributing area of the quarry, estimates of water yield from the quarry were uncertain, so we assumed a uniform specific discharge for the whole LOWER_QUARRY subwatershed.")


### Suspended Sediment Concentration
document.add_heading('Continuous Suspended Sediment Concentration',level=4)
document.add_paragraph("Continuous SSC at 15 minute intervals was estimated from 1) linear interpolation of SSC measured from water samples, and 2) 15 min interval turbidity data (T) and a T-SSC relationship calibrated to stream water samples collected over a range of Q and SSC.")

## Number of SSC samples
def No_All_samples(location):
    No_samples = len(SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin([location])])
    return No_samples
    
## Mean and Max SSC  numbers
def Mean_and_Max_SSC(ssc,location):
    mean =  int(ssc[ssc['Location'].isin([location])]['SSC (mg/L)'].mean())
    maximum = int(ssc[ssc['Location'].isin([location])]['SSC (mg/L)'].max())
    return "{:,}".format(mean), "{:,}".format(maximum)
    
Mean_SSC_FG1, Max_SSC_FG1  = Mean_and_Max_SSC(SSC_dict['Pre-ALL'],'DAM')
Mean_SSC_FG2, Max_SSC_FG2 = Mean_and_Max_SSC(SSC_dict['Pre-ALL'],'DT')
Mean_SSC_FG3, Max_SSC_FG3 = Mean_and_Max_SSC(SSC_dict['Pre-ALL'],'LBJ')

Max_SSC_FG3_Event = "{:%m/%d/%Y}".format(SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['LBJ'])]['SSC (mg/L)'].idxmax())
time_of_max_SSC_FG3 = RoundTo15(SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['LBJ'])]['SSC (mg/L)'].idxmax())
Max_SSC_FG3_Q = "{:,}".format(int(LBJ['Q'][time_of_max_SSC_FG3]))

Max_SSC_FG1_Event = "{:%m/%d/%Y}".format(SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['DAM'])]['SSC (mg/L)'].idxmax())
time_of_max_SSC_FG1 =  RoundTo15(SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['DAM'])]['SSC (mg/L)'].idxmax())
Max_SSC_FG1_Q = "{:,}".format(int(DAM['Q'][time_of_max_SSC_FG1]))

document.add_paragraph("Stream water samples were collected by grab sampling with 500 mL HDPE bottles at FG1, FG2, and FG3. At FG2, water samples were also collected at 30 min intervals during storm events by an ISCO 3700 Autosampler triggered by a stage height sensor. Samples were analyzed for SSC on-island using gravimetric methods (Gray, 2014; Gray et al., 2000). Water samples were vacuum filtered on pre-weighed 47mm diameter, 0.7 um Millipore AP40 glass fiber filters, oven dried at 100 C for one hour, cooled and weighed to determine SSC (mg/L). From January 6, 2012, to October 1, 2014, "+str(len(SSC_dict['Pre-ALL']))+" water samples were collected and analyzed for SSC: FG1 (n="+"%.0f"%No_All_samples('DAM')+"), FG2 (n="+"%.0f"%No_All_samples('DT')+" grab samples, n="+"%.0f"%No_All_samples('R2')+" from the Autosampler), and FG3 (n="+"%.0f"%No_All_samples('LBJ')+").")

document.add_heading('Interpolated grab samples',level=5)
document.add_paragraph("Interpolation of SSC values from grab samples could only be performed if at least three stream water samples were collected during a storm event (Nearing et al., 2007), and if they adequately captured the SSC dynamics of the storm event. SSC was assumed to be zero at the beginning and end of each storm if no grab sample data was available for those times (Lewis et al., 2001).")

document.add_heading('Turbidity-SSC relationships',level=5)
document.add_paragraph("Turbidity (T) was measured at FG1 and FG3 using three types of turbidimeters: 1) Greenspan TS3000 (TS), 2) YSI 600OMS with 6136 turbidity probe (YSI), and 3) Campbell Scientific OBS500 (OBS). All turbidimeters were permanently installed in protective PVC housings near the streambed where the turbidity probe would be submerged at all flow conditions, with the turbidity probe oriented downstream. Despite regular maintenance, debris fouling during storm and baseflows was common and caused data loss during several storm events (Lewis et al., 2001). Storm events with incomplete or invalid T data were not used in the analysis. A three-point calibration was performed on the YSI turbidimeter with YSI turbidity standards (0, 126, and 1000 NTU) at the beginning of each field season and approximately every 3-6 months during data collection. Turbidity measured with 0, 126, and 1000 NTU standards differed by less than 10% (4-8%) during each recalibration. The OBS requires calibration every two years, so recalibration was not needed during the study period. All turbidimeters were cleaned following storms to ensure proper operation.")

document.add_paragraph("At FG3, a YSI turbidimeter recorded T (NTU) at 5 min intervals from January 30, 2012, to February 20, 2012, and at 15 min intervals from February 27, 2012 to May 23, 2012, when it was damaged during a large storm. The YSI turbidimeter was replaced with an OBS, which recorded Backscatter (BS) and Sidescatter (SS) at 5 min intervals from March 7, 2013, to July 15, 2014 (OBSa), and was resampled to 15 min intervals. No data was recorded from August 2013-January 2014 when the wiper clogged with sediment. A new OBS was installed at FG3 from January, 2014, to August, 2014 (OBSb). To correct for some periods of high noise observed in the BS and SS data recorded by the OBSa in 2013, the OBSb installed in 2014 was programmed to make a burst of 100 BS and SS measurements at 15 min intervals, and record Median, Mean, STD, Min, and Max. All BS and SS parameters were analyzed to determine which showed the best relationship with SSC.  Mean SS showed the highest r"u"\u00B2"" and is a physically comparable measurement to NTU measured by the YSI and TS (Anderson, 2005).")

document.add_paragraph("At FG1, the TS turbidimeter recorded T (NTU) at 5 min intervals from January 2012 until it was vandalized and destroyed in July 2012. The YSI turbidimeter, previously deployed at FG3 in 2012, was repaired and redeployed at FG1 and recorded T (NTU) at 5 min intervals from June 2013 to October 2013, and January 2014 to August 2014. T data was resampled to 15 min intervals to compare with SSC samples for the T-SSC relationship, and to correspond to Q for calculating SSY.")

document.add_paragraph("The T-SSC relationship can be unique to each region, stream, instrument or even each storm event (Lewis et al., 2001), and can be influenced by water color, dissolved solids and organic matter, temperature, and the shape, size, and composition of sediment. However, T has proved to be a robust surrogate measure of SSC in streams (Gippel, 1995), and is most accurate when a unique T-SSC relationship is developed for each instrument separately, using in situ grab samples under storm conditions (Lewis, 1996). A unique T-SSC relationship was developed for each turbidimeter, at each location, using 15 min interval T data and SSC samples from storm periods only (Figure "+T_SSC_Rating_Curves['fig_num']+"). A "+'"synthetic"'+" T-SSC relationship was also developed by placing the turbidimeter in a black tub with water, and sampling T and SSC as sediment was added (Appendix 4, Figure 1), but results were not comparable to T-SSC relationships developed under actual storm conditions and were not used in further analyses.")


## LBJ and DAM YSI T-SSC rating curves
if 'T_SSC_Rating_Curves' in locals():
    document.add_picture(T_SSC_Rating_Curves['filename']+'.png',width=Inches(6))
    add_figure_caption(T_SSC_Rating_Curves['fig_num'],"Turbidity-Suspended Sediment Concentration relationships for a) the YSI turbidimeter deployed at FG3 ("+"{:%m/%d/%Y}".format(LBJ_YSI.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(LBJ_YSI.dropna().index[-1])+") and the same YSI turbidimeter deployed at FG1 ("+"{:%m/%d/%Y}".format(DAM_YSI.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(DAM_YSI.dropna().index[-1])+"). b) OBS500 turbidimeter deployed at FG3 ("+"{:%m/%d/%Y}".format(LBJ_OBSa.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(LBJ_OBSa.dropna().index[-1])+") and c) OBS500 turbidimeter deployed at FG3 ("+"{:%m/%d/%Y}".format(LBJ_OBSb.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(LBJ_OBSb.dropna().index[-1])+").")


### Maybe use the LBJ_YSI T-SSC for FG1 also; same arguement as for the TS3K (low range of T-SSC
document.add_paragraph("The T-SSC relationships varied among sampling sites and sensors but all showed acceptable r2 values ("+"%.2f"%LBJ_YSI_rating.r2+"-"+"%.2f"%DAM_YSI_rating.r2+"). Lower scatter was achieved by using grab samples collected during stormflows only. For the TS (not shown) and YSI deployed at FG1, the r2 values were high ("+"%.2f"%DAM_TS3K_rating.r2+", "+"%.2f"%DAM_YSI_rating.r2+") but the ranges of T and SSC values used to develop the relationships were considered too small (0-"+"%.0f"%T_SSC_DAM_TS3K[1]['T-NTU'].max()+" NTU) compared to the maximum observed during the deployment period ("+"{:,}".format(int(DAM_TS3K['NTU'].max()))+" NTU) to develop a robust relationship for higher T values. Instead, the T-SSC relationship developed for the YSI turbidimeter installed at FG3 (Figure "+T_SSC_Rating_Curves['fig_num']+"a) was used to calculate SSC from T data collected by the TS and the YSI at FG1. For the YSI turbidimeter, more scatter was observed in the T-SSC relationship at FG3 than at FG1 (Figure "+T_SSC_Rating_Curves['fig_num']+"a), which could be attributed to the higher number and wider range of values sampled, and to temporal variability in sediment characteristics. The OBSa and OBSb turbidimeters had high r2 values ("+"%.2f"%LBJ_OBSa_rating.r2+", "+"%.2f"%LBJ_OBSb_rating.r2+") and compared well between the two periods of deployment (Figure "+T_SSC_Rating_Curves['fig_num']+"b).")

### Cumulative Probable Error
document.add_heading('Cumulative Probable Error (PE)',level=3)

## Duvert 2010 writes Q error is 10-20% and some other errors, should incorporate his stuff
document.add_paragraph("Uncertainty in SSYEV estimates arises from both measurement and model errors, including models of stage-discharge (stage-Q) and turbidity-suspended sediment concentration (T-SSC) (Harmel et al., 2006). The Root Mean Square Error (RMSE) method estimates the "+'"most probable value"'+" of the cumulative or combined error by propagating the error from each measurement and modeling procedure to the final SSYEV calculation (Topping, 1972). The resulting cumulative probable error (PE) is the square root of the sum of the squares of the maximum values of the separate errors:")

add_equation(PE_eq) ## Equation

document.add_paragraph("EQmeas  and ESSCmeas were estimated using lookup tables (LUT) from the DUET-H/WQ software tool (Harmel et al., 2009). The effect of uncertain SSYEV estimates may complicate conclusions about contributions from subwatersheds, anthropogenic impacts, and SSYEV-Storm Metric relationships. This is common in sediment yield studies where successful models estimate SSY with "u"\u00B1""50-100% accuracy (Duvert et al., 2012) but the difference in SSY from undisturbed and disturbed areas was expected to be much larger than the cumulative uncertainty. PE was calculated for SSYEV from the UPPER and TOTAL watersheds, but not calculated for SSYEV from the LOWER subwatershed since it was calculated as the difference of SSYUPPER and SSYTOTAL.") 

#### RESULTS ####
results_title = document.add_heading('Results',level=2)
## Field data collection
document.add_heading('Field Data Collection',level=3)

#### Precipitation
document.add_heading('Precipitation',level=4)
precip_2012 = "{:,}".format(int(PrecipFilled[start2012:stop2012].sum()))
precip_2013 = "{:,}".format(int(PrecipFilled[start2013:stop2013].sum()))
precip_2014 = "{:,}".format(int(PrecipFilled[start2014:stop2014].sum()))
percent_of_annual_mean_P = (PrecipFilled['Precip'][start2012:stop2012].sum() + PrecipFilled['Precip'][start2013:stop2013].sum() + PrecipFilled['Precip'][start2014:dt.datetime(2014,12,31)].sum()) / (3 * 3800) * 100

document.add_paragraph("Annual precipitation measured at RG1, with gaps filled with data from Wx, was "+precip_2012+" mm, "+precip_2013+" mm, and "+precip_2014+" mm in 2012, 2013, and 2014, respectively, which are approximately "+"%.0f"%percent_of_annual_mean_P+"% of long-term precipitation (=3,800 mm) from PRISM data (Craig, 2009). No difference in measured P was found between RG1 and Wx, or between RG1 and RG2, so P was assumed to be homogenous over the watershed for all analyses. Rain gauges could only be placed as high as ~300 m (RG2), though the highest point in the watershed is ~600 m. Long-term rain gage records show a strong precipitation gradient with increasing elevation, with average precipitation of 3,000-4,000 mm on the lowlands, increasing to more than 6,350 mm at high elevations (>400 m.a.s.l.) (Craig, 2009; Dames & Moore, 1981; Wong, 1996). Precipitation data measured at higher elevations would be useful to determine a more robust orographic precipitation relationship. For this analysis, however, the absolute values of total precipitation in each subwatershed are not as important since precipitation and the erosivity index are only used as predictive storm metrics.")

#### Water Discharge
document.add_heading('Water Discharge',level=4)

document.add_paragraph("Discharge at both FG1 and FG3 was characterized by periods of low but perennial baseflow, punctuated by short, flashy hydrograph peaks (FG1: max "+"{:,g}".format(DAM['Q'].max())+" L/sec, FG3: max "+"{:,g}".format(LBJ['Q'].max())+" L/sec) (Figure "+Q_timeseries['fig_num']+"). Though Q data was unavailable for some periods, storm events appeared to be generally smaller but more frequent in the October-April wet season. Storm events during the May-September dry season were less frequent but larger. The largest event in the three year study was observed in August 2014.")

if 'Q_timeseries' in locals():
    document.add_picture(Q_timeseries['filename']+'.png',width=Inches(6))
    add_figure_caption(Q_timeseries['fig_num'],"Time series of water discharge (Q), calculated from measured stage and the stage-discharge rating curves in a) 2012 b) 2013 and c) 2014.")

#### Number of Storm Events
document.add_heading('Storm Events',level=4)
## All Storms identified from Q at FG1 and FG3
No_of_Storm_Intervals = len(All_Storms[All_Storms['start']<Mitigation])
## Storms with P and SSY data
No_of_Storm_Intervals_DAM_S = len(Storms_DAM[Storms_DAM['Sstart']<Mitigation]['Ssum'].dropna())
No_of_Storm_Intervals_LBJ_S = len(Storms_LBJ[Storms_LBJ['Sstart']<Mitigation]['Ssum'].dropna())
No_of_Storm_Intervals_QUA_S = S_budget_2['Storm Start']['Total/Avg']
## Max storm duration
max_storm_duration = All_Storms['duration (hrs)'].max()/24

document.add_paragraph("Using the storm definition criteria, "+"%.0f"%No_of_Storm_Intervals+" events were identified from Q data at FG1 and FG3 between January, 2012, to July 2014; "+str(len(Q_budget)-1)+" events had simultaneous Q data at FG1 and FG3 (Appendix 3, Table "+Q_budget_table.table_num+"). SSC data from T or interpolated grab samples were recorded during "+"%.0f"%No_of_Storm_Intervals_DAM_S+" events at FG1, and "+"%.0f"%No_of_Storm_Intervals_LBJ_S+" events at FG3. Of those storms, "+S_budget['Storm Start']['Total/Avg']+" events had data for P, Q, and SSC at both FG1 and FG3 to calculate SSY from the LOWER subwatershed. SSY data from interpolated grab samples were collected at FG2 for "+No_of_Storm_Intervals_QUA_S+" storms to calculate SSY from the LOWER_QUARRY and LOWER_VILLAGE subwatersheds separately. Storm event durations ranged from "+"%.0f"%All_Storms['duration (hrs)'].min()+" hours to "+"%.0f"%max_storm_duration+" days, with mean duration of "+"%.0f"%All_Storms['duration (hrs)'].mean()+" hours.")


document.add_paragraph("Most storm events showed a typical pattern, where a short period of intense rainfall caused a rapid increase in SSC downstream of the quarry (FG2) while SSC remained low at the undisturbed forest site (FG1)(Figure "+Example_Storm['fig_num']+"). The highest SSC was typically observed at FG2, with slightly lower and later peak SSC observed at FG3. SSC downstream of the undisturbed forest (FG1) typically increased more slowly, remained lower, and peaked later than the disturbed sites downstream of the quarry (FG2) and the village (FG3). Though peak SSC was highest at FG2, the highest SSY was measured at FG3 due to the addition of storm runoff from the larger watershed draining to FG3.")

if 'Example_Storm' in locals():
    document.add_picture(Example_Storm['filename']+'.png',width=Inches(6))
    add_figure_caption(Example_Storm['fig_num'],"Example of storm event ("+"{:%m/%d/%Y}".format(example_storm_start)+"). SSY at FG1 and FG3 calculated from SSC modeled from T, and SSY at FG2 from SSC samples collected by the Autosampler.")
    
#### Suspended Sediment Concentration
document.add_heading('Suspended Sediment Concentration',level=4)

def Mean_stormflow_SSC(location):    
    Mean_storm_samples  = SSC_dict['Pre-storm'][SSC_dict['Pre-storm']['Location'].isin([location])]['SSC (mg/L)'].mean()
    return "%.0f"%Mean_storm_samples 
def No_storm_samples(location):
    No_storm = len(SSC_dict['Pre-storm'][SSC_dict['Pre-storm']['Location'].isin([location])])
    return "%.0f"%No_storm
def Percent_storm_samples(location):
    storm_samples = len(SSC_dict['Pre-storm'][SSC_dict['Pre-storm']['Location'].isin([location])])
    all_samples = len(SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin([location])]) 
    percent_storm = storm_samples/all_samples *100
    return "%.0f"%percent_storm
Percent_stormflow_FG1 = Percent_storm_samples('DAM')
Percent_stormflow_FG2 = Percent_storm_samples('DT')
Percent_stormflow_FG3 = Percent_storm_samples('LBJ')

def Mean_nonstormflow_SSC(location):    
    Mean_nonstormflow_samples  = SSC_dict['Pre-baseflow'][SSC_dict['Pre-baseflow']['Location'].isin([location])]['SSC (mg/L)'].mean()
    return "%.0f"%Mean_nonstormflow_samples 
def No_nonstormflow_samples(location):
    No_nonstormflow = len(SSC_dict['Pre-baseflow'][SSC_dict['Pre-baseflow']['Location'].isin([location])])
    return "%.0f"%No_nonstormflow
def Percent_nonstormflow_samples(location):
    nonstormflow_samples = len(SSC_dict['Pre-baseflow'][SSC_dict['Pre-baseflow']['Location'].isin([location])])
    all_samples = len(SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin([location])]) 
    percent_nonstormflow = nonstormflow_samples/all_samples *100
    return "%.0f"%percent_nonstormflow
Percent_nonstormflow_FG1 = Percent_nonstormflow_samples('DAM')
Percent_nonstormflow_FG2 = Percent_nonstormflow_samples('DT')
Percent_nonstormflow_FG3 = Percent_nonstormflow_samples('LBJ')

f1, p1, f2, p2 = "%.3f"%f1, "%.3f"%p1, "%.3f"%f2, "%.3f"%p2
QUARRY_DAM_ttest1_p, QUARRY_DAM_ttest2_p = "%.3f"%QUARRY_DAM_ttest1[1], "%.3f"%QUARRY_DAM_ttest2[1]
QUARRY_LBJ_ttest1_p, QUARRY_LBJ_ttest2_p = "%.3f"%QUARRY_LBJ_ttest1[1], "%.3f"%QUARRY_LBJ_ttest2[1]

H1, KWp1, H2, KWp2 = "%.3f"%H1, "%.3f"%KWp1, "%.3f"%H2, "%.3f"%KWp2
QUARRY_DAM_mannwhit1_p, QUARRY_DAM_mannwhit2_p = "%.3f"%(QUARRY_DAM_mannwhit1[1]*2), "%.3f"%(QUARRY_DAM_mannwhit2[1]*2)
QUARRY_LBJ_mannwhit1_p, QUARRY_LBJ_mannwhit2_p = "%.3f"%(QUARRY_LBJ_mannwhit1[1]*2), "%.3f"%(QUARRY_LBJ_mannwhit2[1]*2)

## SSC boxplots
document.add_paragraph("Mean ("u"\u03BC"") and maximum SSC of water samples, collected during non-stormflow and stormflow periods by grab and autosampler, were lowest at FG1 ("u"\u03BC""="+Mean_SSC_FG1+" mg/L, max="+Max_SSC_FG1+" mg/L), highest at FG2 ("u"\u03BC""="+Mean_SSC_FG2+" mg/L, max="+Max_SSC_FG2+" mg/L), and in between at FG3 ("u"\u03BC""="+Mean_SSC_FG3+" mg/L, max="+Max_SSC_FG3+" mg/L). At FG1, "+Percent_nonstormflow_FG1+"% of grab samples (n="+No_nonstormflow_samples('DAM')+") were collected during non-stormflow ("u"\u03BC""="+Mean_nonstormflow_SSC('DAM')+" mg/L (Figure "+SSC_Boxplots['fig_num']+"a); "+Percent_stormflow_FG1+"% of grab samples (n="+No_storm_samples('DAM')+") were collected during stormflow, "u"\u03BC""= "+Mean_stormflow_SSC('DAM')+" mg/L (Figure 8b). At FG2, "+Percent_nonstormflow_FG2+"% of grab samples (n="+No_nonstormflow_samples('DT')+") were collected during non-stormflow ("u"\u03BC""= "+Mean_nonstormflow_SSC('DT')+" mg/L; "+Percent_stormflow_FG2+"% of grab samples (n="+No_storm_samples('DT')+") were collected during stormflow, "u"\u03BC""= "+Mean_stormflow_SSC('DT')+" mg/L. At FG3, "+Percent_nonstormflow_FG3 +"% of samples (n="+No_nonstormflow_samples('LBJ')+") were collected during non-stormflow ("u"\u03BC""= "+Mean_nonstormflow_SSC('LBJ')+" mg/L; "+Percent_stormflow_FG3+"% of samples (n="+No_storm_samples('LBJ')+") were collected during stormflow, "u"\u03BC""= "+Mean_stormflow_SSC('LBJ')+" mg/L. This pattern of SSC values suggests that little sediment is contributed from the forest upstream of FG1, followed by a large input of sediment between FG1 and FG2, and then SSC is diluted by addition of stormflow with lower SSC between FG2 and FG3.")

## Mean SSC statistical tests
### ARE THESE RIGHT??
document.add_paragraph("Probability plots of the SSC data collected at FG1, FG2 and FG3 showed they were highly non-normal, so non-parametric tests for statistical significance were applied. The Kruskall-Wallis test showed SSC samples from all three locations were significantly different for non-stormflow (p<"+KWp1+") and stormflow (p<"+KWp1+"). The pair-wise Mann-Whitney test showed SSC samples were significantly different between FG1 and FG2 (non-stormflow, p="+QUARRY_DAM_mannwhit1_p+"; stormflow, p="+QUARRY_DAM_mannwhit2_p+"), but were not significantly different between FG2 and FG3 (non-stormflow, p="+QUARRY_LBJ_mannwhit1_p+"; stormflow, p="+QUARRY_LBJ_mannwhit2_p+").") 

## Figure SSC_Boxplots
if 'SSC_Boxplots' in locals():
    document.add_picture(SSC_Boxplots['filename']+'.png',width=Inches(6))
    add_figure_caption(SSC_Boxplots['fig_num'],"Boxplots of Suspended Sediment Concentration (SSC) from grab samples only (no Autosampler) at FG1, FG2, and FG3 during (a) non-stormflow and (b) stormflow.")
    

## Water Discharge vs Sediment Concentration  
document.add_paragraph("SSC varied by several orders of magnitude for a given Q at FG1, FG2, and FG3 due to significant hysteresis observed during storm periods (Figure "+Discharge_Concentration['fig_num']+"). At FG1, interannual variability of SSC during stormflow was assumed to be caused by randomly occurring landslides or mobilization of sediment stored in the watershed during large storm events. The maximum SSC sampled downstream of the undisturbed forest, at FG1 ("+Max_SSC_FG1+" mg/L), was sampled on "+Max_SSC_FG1_Event+" at high discharge (QFG1= "+Max_SSC_FG1_Q+" L/sec) (Figure "+Discharge_Concentration['fig_num']+"a). Anecdotal and field observations reported higher than normal SSC upstream of the quarry during the 2013 field season, possibly due to landsliding from previous large storms (G. Poysky, pers. comm.).")

## Figure Water Discharge vs Sediment Concentration
if 'Discharge_Concentration' in locals():
    document.add_picture(Discharge_Concentration['filename']+'.png',width=Inches(6))
    add_figure_caption(Discharge_Concentration['fig_num'],"Water Discharge vs Suspended Sediment Concentration at a) FG1, b) FG2, and c) FG3 during non-stormflow and stormflow periods. The box in b) highlights the samples with high SSC during low flows, solid symbols indicate SSC samples where precipitation during the preceding 24 hours was 0 mm.") 


document.add_paragraph("At FG2 and FG3, additional variability in the Q-SSC relationship was due to the changing sediment availability associated with quarrying operations and construction in the village. The high SSC values observed downstream of the quarry (FG2) during low Q were caused by two mechanisms: 1) precipitation events that did not result in stormflow, but generated runoff from the quarry with high SSC and 2) washing fine sediment into the stream during rock crushing operations at the quarry. ")

document.add_paragraph("The maximum SSC sampled at FG2 ("+Max_SSC_FG2+" mg/L) and FG3 ("+Max_SSC_FG3+" mg/L) were sampled during the same rainfall event ("+Max_SSC_FG3_Event+"), but during low Q (QFG3="+Max_SSC_FG3_Q+" L/sec)(Figure "+Discharge_Concentration['fig_num']+"b-c). During this event, brief but intense precipitation caused high sediment runoff from the quarry, but did not increase Q above the defined storm threshold. SSC was diluted further downstream of the quarry at FG3 by the addition of runoff with lower SSC from the village.")

document.add_paragraph("Given the close proximity of the quarry to the stream, SSC downstream of the quarry can be highly influenced by mining activity like rock extraction, crushing, and/or hauling operations. During 2012, a common practice for removing fine sediment from crushed aggregate was to rinse it with water pumped from the stream. In the absence of retention structures the fine sediment was then discharged directly into the stream, causing high SSC during baseflow periods with no precipitation in the preceding 24 hours (solid symbols, Figure "+Discharge_Concentration['fig_num']+"b-c). Riverine discharge of fine sediment rinsed from aggregate was discontinued in 2013, corresponding with low SSC during low Q in 2013 (Figure "+Discharge_Concentration['fig_num']+"b-c). In 2013 and 2014, waste sediment was piled on-site and severe erosion of these changing stockpiles caused high SSC during storm events.")


#### Cumulative Probable Error
document.add_heading('Cumulative Probable Error (PE)',level=4)
## PE for Water Discharge Q
document.add_paragraph("The measurement error (RMSE) was "+"%.1f"%AV_Q_measurement_RMSE+" % for Q at FG1 and FG3 from the DUET-H/WQ LUT (Harmel et al., 2006), which included error in the area-velocity measurements (6%), continuous Q measurement in a natural channel (6%), pressure transducer error (0.1%), and streambed condition (firm, stable bed=0%). The model errors (RMSE) were "+"%.0f"%LBJ_Man_rmse+"% for the stage-Q rating curve using Manning's equation at FG3, and "+"%.0f"%DAM_HEC_rmse+"% using HEC-RAS at FG1.")
## PE for SSC and SSY
document.add_paragraph("The measurement errors (RMSE) for SSC measurements from the DUET/WQ LUT, were 16.3% for sample collection, which included error from interpolating over a 30 min interval (5%), sampling during stormflows (3%), and 3.9% for sample analysis which included measuring SSC by filtration (3.9%). The model errors (RMSE) of the T-SSC relationships were "+"%.1f"%DAM_YSI_rating_rmse+"% ("+"%.0f"%T_SSC_DAM_YSI[0].rmse+" mg/L) for the YSI and TS at FG1, "+"%.1f"%LBJ_YSI_rating_rmse+"% ("+"%.0f"%T_SSC_LBJ_YSI[0].rmse+" mg/L) for the YSI at FG3, and "+"%.1f"%LBJ_OBS_rating_rmse+"% ("+"%.0f"%T_SSC_LBJ_OBS[0].rmse+" mg/L) for the OBS at FG3. Cumulative Probable Error (RMSE %) for SSY estimates at FG1 and FG3 were calculated from the measurement errors for Q (8.5%) and SSC grab samples (16.3%), and the model errors of the respective stage-Q and T-SSC relationships for that location. Cumulative Probable Errors (PE) in SSYEV were "+"%.0f"%S_budget['UPPER PE %'].min()+"-"+"%.0f"%S_budget['UPPER PE %'][:-3].max()+"% ("u"\u03BC""="+S_budget['UPPER PE %']['Total/Avg']+"%) at FG1 and "+"%.0f"%S_budget['TOTAL PE %'].min()+"-"+"%.0f"%S_budget['TOTAL PE %'][:-3].max()+"% ("u"\u03BC""="+S_budget['TOTAL PE %']['Total/Avg']+"%) at FG3. ")


#### Comparing SSYEV from disturbed and undisturbed subwatersheds
document.add_heading('Comparing SSYEV from disturbed and undisturbed subwatersheds',level=3)

## These are PE in tons for the plus-minus
PE_SSY_TOTAL_2 = float(S_budget['TOTAL tons']['Total/Avg']) * (float(S_budget['TOTAL PE %']['Total/Avg']) / 100)
PE_sSSY_TOTAL_2 = float(S_budget['TOTAL tons']['Tons/km2']) * (float(S_budget['TOTAL PE %']['Total/Avg']) / 100)
## These are PE in tons for the plus-minus
PE_SSY_UPPER_2 = float(S_budget['UPPER tons']['Total/Avg']) * (float(S_budget['UPPER PE %']['Total/Avg']) / 100)
PE_sSSY_UPPER_2 = float(S_budget['UPPER tons']['Tons/km2']) * (float(S_budget['UPPER PE %']['Total/Avg']) / 100)


## Storm Sediment Budget Table: UPPER, LOWER, TOTAL
document.add_paragraph("SSYEV was measured simultaneously at FG1 and FG3 for "+S_budget['Storm Start']['Total/Avg']+" storms (Table "+S_budget_table.table_num+"). SSYTOTAL was "+SSY_TOTAL_2+u"\u00B1"+"%.1f"%PE_SSY_TOTAL_2+" tons ("+sSSY_TOTAL_2+u"\u00B1"+"%.1f"%PE_sSSY_TOTAL_2+" tons/km"u"\u00B2""), with "+SSY_UPPER_2+u"\u00B1"+"%.1f"%PE_SSY_UPPER_2+" tons ("+sSSY_UPPER_2+u"\u00B1"+"%.1f"%PE_sSSY_UPPER_2+" tons/km"u"\u00B2"") from the UPPER subwatershed and "+SSY_LOWER_2+" tons ("+sSSY_LOWER_2+" tons/km"u"\u00B2"") from the LOWER subwatershed. The UPPER and LOWER subwatersheds are similar in size (0.90 km"u"\u00B2"" and 0.88 km"u"\u00B2"") but SSYUPPER accounted for an average of just "+Percent_Upper_S_2 +"% and SSYLOWER for "+Percent_Lower_S_2 +"% of SSY TOTAL at the watershed outlet (Table "+S_budget_table.table_num+"). The DR estimated from sSSYUPPER and sSSYLOWER suggests sSSY has increased by "+S_budget['LOWER tons']['DR']+"x in the LOWER subwatershed, and "+S_budget['TOTAL tons']['DR']+"x for the TOTAL watershed.")


document.add_paragraph("The measured sSSY from the forested UPPER watershed (sSSYUPPER="+sSSY_UPPER_2+" tons/km"u"\u00B2"") was used to calculate SSY from undisturbed forest in the LOWER subwatershed. SSY from the undisturbed forest areas in the LOWER watershed was "+S_budget_analysis['LOWER']['Forested areas']+" tons, so SSY from the disturbed areas was "+S_budget_analysis['LOWER']['Disturbed areas']+" tons (Equation "+SSY_disturbed.eq_num+"). For the storms in Table "+S_budget_table.table_num+", roughly "+S_budget_analysis['LOWER']['percent from disturbed areas']+"% of SSYLOWER was from disturbed areas, despite the disturbed areas only accounting for "+S_budget_analysis['LOWER']['Fraction of subwatershed area disturbed']+"% of the LOWER subwatershed area (0.089 km"u"\u00B2""). Similarly, despite only "+S_budget_analysis['TOTAL']['Fraction of subwatershed area disturbed']+"% of the TOTAL watershed being disturbed, SSY from disturbed areas accounted for "+S_budget_analysis['TOTAL']['percent from disturbed areas']+"% of the SSYTOTAL. sSSY from disturbed areas in the LOWER subwatershed was "+"{:,g}".format(float(S_budget_analysis['LOWER']['sSSY, disturbed areas (tons/km2)']))+" tons/km"u"\u00B2"", or "+S_budget_analysis['LOWER']['DR for sSSY from disturbed areas']+"x the sSSY of undisturbed forest.")



## Storm Sediment Budget Table: UPPER, LOWER, LOWER_QUARRY, LOWER_VILLAGE, TOTAL


document.add_paragraph("SSYEV was measured simultaneously at FG1, FG2, and FG3 for "+No_of_Storm_Intervals_QUA_S+" of the storms in Table "+S_budget_table.table_num+", so SSYEV from the LOWER subwatershed containing the quarry (SSYLOWER_QUARRY) and LOWER subwatershed containing the village below the quarry (SSYLOWER_VILLAGE) could be calculated separately (Table "+S_budget_2_table.table_num+").")

#document.add_paragraph("Storm on 2/3/12 has a potential outlier at DT at beginning of storm, which makes the SSYquarry huge! Storm at 2/5/12 doesn't have adequate SSC samples for quarry and its a multipeaked event so the SSY doesn't fall back to low levels like the LBJ and DAM T data suggest it should. Storm 3/6/13 looks like may have missed the second peak but the data looks comparable between sites. Storm 4/16/13 doesn't have alot of points but maybe they're adequate? Storm 4/23/13 looks good. Storm 4/30/13 has inadequate SSC data for all locations after the first peak; can maybe change the storm interval? Storm 6/5/13 has inadequate SSC data for all locations for the first peak, decent data for LBJ and DT for second peak but not for DAM; can maybe change the storm interval? Storm 2/14/14 looks good. Storm 2/20/14 looks good. Storm 2/21/14 looks good. Storm 2/27/14 looks kinda shitty. So good storms are: 3/6/13, 4/16/13? 4/23/13, 4/30/13?, 6/5/13?, 2/14/14, 2/20/14, 2/21/14") 

document.add_paragraph("For the "+No_of_Storm_Intervals_QUA_S+" storms in Table "+S_budget_2_table.table_num+", SSYTOTAL was "+SSY_TOTAL_4+" tons with an average of "+Percent_UPPER_S_4+"% from the UPPER subwatershed, "+Percent_QUARRY_S+"% from LOWER_QUARRY subwatershed, and "+Percent_VILLAGE_S+"% from the LOWER_VILLAGE subwatershed. sSSY from the UPPER, LOWER_QUARRY, and LOWER_VILLAGE subwatersheds, and the TOTAL watershed was "+sSSY_UPPER_4+", "+sSSY_LOWER_QUARRY+", "+sSSY_LOWER_VILLAGE+", and "+sSSY_TOTAL_4+" tons/km"u"\u00B2"", respectively. sSSY from LOWER_QUARRY and LOWER_VILLAGE subwatersheds was "+S_budget_2_analysis['LOWER_QUARRY']['DR for sSSY from disturbed areas']+"x and "+S_budget_2_analysis['LOWER_VILLAGE']['DR for sSSY from disturbed areas']+"x higher, respectively, than sSSY from UPPER subwatershed, suggesting human disturbance has significantly increased SSY over natural levels, particularly at the quarry. sSSYTOTAL was "+S_budget_2['TOTAL tons']['DR']+"x higher than the sSSYUPPER, similar to the larger range of storms in Table 2, where sSSY was "+S_budget['TOTAL tons']['DR']+"x higher than undisturbed forest conditions.")

document.add_paragraph("Very small fractions of the subwatershed areas are disturbed, yet roughly "+ S_budget_2_analysis['LOWER_QUARRY']['percent from disturbed areas']+"% of SSYLOWER_QUARRY ("+percent_disturbed_area_LOWER_QUARRY+"% disturbed) and "+ S_budget_2_analysis['LOWER_VILLAGE']['percent from disturbed areas']+"% of SSYLOWER_VILLAGE ("+percent_disturbed_area_LOWER_VILLAGE+"% disturbed) subwatersheds was from disturbed areas. Similarly, despite only "+S_budget_analysis['TOTAL']['Fraction of subwatershed area disturbed']+"% of the TOTAL watershed being disturbed, "+ S_budget_analysis['TOTAL']['percent from disturbed areas']+"-"+S_budget_2_analysis['TOTAL']['percent from disturbed areas']+"% of SSYTOTAL (Tables "+S_budget_analysis_table.table_num+" and "+S_budget_2_analysis_table.table_num+") was from disturbed areas. Bare land in the LOWER_QUARRY subwatershed significantly increased sSSYLOWER_QUARRY and sSSYTOTAL, and contributed the majority of SSY from disturbed areas in Faga'alu watershed. sSSY from disturbed areas in the UPPER, LOWER_QUARRY, and LOWER_VILLAGE subwatersheds was "+S_budget_2_analysis['UPPER']['sSSY, disturbed areas (tons/km2)']+", "+"{:,g}".format(float(S_budget_2_analysis['LOWER_QUARRY']['sSSY, disturbed areas (tons/km2)']))+", and "+S_budget_2_analysis['LOWER_VILLAGE']['sSSY, disturbed areas (tons/km2)']+" tons/km"u"\u00B2"", respectively, suggesting that disturbed areas increase sSSY over forested conditions by "+S_budget_2_analysis['LOWER_QUARRY']['DR for sSSY from disturbed areas']+"x and "+S_budget_2_analysis['LOWER_VILLAGE']['DR for sSSY from disturbed areas']+"x in the LOWER_QUARRY and LOWER_VILLAGE subwatersheds, respectively. Human disturbance in the LOWER_VILLAGE subwatershed also increased SSY above natural levels but the magnitude of disturbance was much lower than the quarry.")

    
#### Predicting SSYEV from storm metrics
document.add_heading('Predicting SSYEV from storm metrics',level=3)

document.add_paragraph("SSYEV from the UPPER and TOTAL watersheds correlated with each of the four storm metrics tested (Figure "+SSY_models_ALL['fig_num']+"). Precipitation metrics (Psum and EI30) showed lower Pearson and Spearman correlation coefficients compared to the discharge metrics (Qsum and Qmax) (Table "+All_Models_stats_table.table_num+"). SSYEV is calculated from Q so it is expected that discharge metrics are more closely correlated, and this has also been observed in other studies (Duvert et al., 2010; Rankl, 2004). Pearson and Spearman correlation coefficients were fairly similar, meaning the relationships were mostly linear in log-log space. Significant scatter was observed around all models, which reflects the changing sediment availability at the quarry and village, and the natural variability in the watershed response for different storm events.")

document.add_paragraph("Qmax was the best predictor of SSYEV for both the UPPER and TOTAL watersheds. The Qmax model for both UPPER and TOTAL watersheds showed the highest coefficient of determination (r2), lowest RMSE, and highest Pearson and Spearman correlation coefficients (Table "+All_Models_stats_table.table_num+"). Qsum showed an equally high r2, but only for the UPPER subwatershed, and RMSE was higher in both subwatersheds for Qsum than for Qmax. Discharge metrics showed much higher correlation coefficients than the precipitation metrics in the UPPER subwatershed, but were more similar in the LOWER watershed. This suggests that sediment production is more related to discharge processes in the UPPER subwatershed, and more related to precipitation processes in the LOWER subwatershed. Precipitation was measured at the quarry, which may reflect precipitation characteristics more accurately in the LOWER than the UPPER watershed. SSY from the LOWER subwatershed is hypothesized to be mostly generated by hillslope erosion by sheetwash and rill formation at the quarry and on dirt roads, and agricultural plots, whereas SSY from the UPPER subwatershed is hypothesized to be mainly from channel processes and mass wasting. Mass wasting can contribute large pulses of sediment which can be deposited near or in the streams and entrained at high discharges during later storm events. Given the high correlation coefficients between SSYEV and Qmax in both subwatersheds, Qmax may be a promising predictor that integrates both precipitation and discharge processes.")

if 'SSY_models_ALL' in locals():
    document.add_picture(SSY_models_ALL['filename']+'.png',width=Inches(6))
    add_figure_caption(SSY_models_ALL['fig_num'],"SSYEV regression models for predictive storm metrics. Each point represents a different storm event. **=slopes and intercepts were statistically different (p<0.05), *=intercepts were statistically different (p<0.01).")

document.add_paragraph("In all models, SSYEV from the TOTAL watershed was higher than from the UPPER subwatershed for the full range of measured storms with the exception of a few events that are considered outliers. These events could be attributed to measurement error or to landslides in the UPPER subwatershed and the increased sediment supply for that specific event. Storm sequence and antecedent conditions may also play a role. While the climate on Tutuila is tropical, without strong seasonality, periods of low rainfall can persist for several weeks, perhaps altering the water and sediment dynamics in the subsequent storm events.")

## Storm size
document.add_paragraph("ANCOVA was used to compare regression coefficients ("u"\u03B2""=slope and "u"\u03B1""=intercept) of the UPPER and TOTAL SSY models, to determine if the relative sediment contribution from undisturbed and human-disturbed areas changed with storm size. All model intercepts were significantly different (p<0.05), but only the Psum-SSYEV model showed significantly different (p<0.05) slopes. It was hypothesized that for large storms, SSYEV from the UPPER watershed may become relatively more important for SSY at the outlet, however, the models show conflicting results. The Psum-SSYEV models indicate that for larger storm events SSY from the UPPER and TOTAL watersheds are more similar, as the regression lines converge at higher Psum values. Conversely, the Qsum- and Qmax-SSYEV models show no change in relative contributions of SSY over the range of storm sizes (Figure "+SSY_models_ALL['fig_num']+"). In that case, the discharge models (Qsum and Qmax) support the conclusion that human disturbance as a fraction of total SSY does not diminish with storm size, while the Psum model supports the conclusion that human-disturbance as a fraction of total SSY does diminish with storm size.")



### "Annual estimates of SSY and sSSY"
document.add_heading("Annual estimates of SSY and sSSY",level=3)

## From SSYEV-Qmax relationship
document.add_paragraph("The Qmax-SSY relationships were used to predict SSY from Qmax of "+str(len(no_storms_2014))+" storms in 2014, the only year with a continuous Q record (Table "+est_Annual_SSY_table.table_num+"). Predicted annual SSY in 2014 from the UPPER and TOTAL watersheds was "+SSY_Upper_Qmax_2014+" and "+SSY_Total_Qmax_filled_2014+" tons/year, respectively. Predicted annual sSSY in 2014 from the UPPER and TOTAL watersheds, was "+sSSY_Upper_Qmax_2014+" and "+sSSY_Total_Qmax_filled_2014+" tons/km"u"\u00B2""/year, respectively.")

## Annual SSY and sSSY from full range of measured storms at FG1 and FG3
document.add_paragraph("Annual SSY was also calculated using Equation "+SSY_annual_eq.eq_num+" for three sets of storm events: a) all events with SSYEV data, including those where SSYEV data were only available for a single site; b) only events where data was available for both UPPER (FG1) and TOTAL (FG3) and c) only events where data was available for UPPER (FG1), LOWER_QUARRY (FG2), and TOTAL (FG3). Including all storms (method a) will provide the best estimate at a given location, while b) and c) allow more direct comparison of different subwatersheds. Continuous records of Q and precipitation in 2014 showed annual storm precipitation (Psann) was "+"{:,g}".format(int(P_2014_storm))+" mm, representing "+P_2014_perc_ann+"% of total annual precipitation ("+ "{:,g}".format(int(PrecipFilled[start2014:stop2014].sum()))+" mm). All storms with measured SSY at FG1 from 2012-2014 included"+ "{:,g}".format(int(P_FG1_all_storms))+" mm of precipitation (Psmeas), or "+"%.0f"%P_FG1_percent_storm+"% of Psann, so estimated annual SSY from the UPPER subwatershed from Equation "+SSY_annual_eq.eq_num+" was "+"%.0f"%annual_SSY_UPPER_ALL+" tons/yr ("+"%.0f"%annual_sSSY_UPPER_ALL+" tons/km"u"\u00B2""/yr). All storms with measured SSY at FG3 from 2012-2014 included "+"{:,g}".format(int(P_FG3_all_storms))+" mm of precipitation, or "+"%.0f"%P_FG3_percent_storm+"% of expected annual storm precipitation so estimated annual SSY from the TOTAL watershed was "+"%.0f"%annual_SSY_TOTAL_ALL+" tons/yr ("+"%.0f"%annual_sSSY_TOTAL_ALL+" tons/km"u"\u00B2""/yr).")
    

## Annual SSY and sSSY from Table 2 (UPPER and LOWER)
document.add_paragraph("For storms with measured SSY at both FG1 and FG3 (Table "+S_budget_2_table.table_num+") Psmeas was "+"{:,g}".format(int(P_measured_2))+" mm, or "+"%.0f"%P_measured_2_perc_storm+"% of Psann. Using Equation "+SSY_annual_eq.eq_num+", estimated annual SSY from the UPPER, LOWER, and TOTAL watersheds was "+annual_SSY_UPPER_2+", "+annual_SSY_LOWER_2+", and "+annual_SSY_TOTAL_2+" tons/year, respectively. Estimated annual sSSY from the UPPER, LOWER, and TOTAL watersheds was "+annual_sSSY_UPPER_2+", "+annual_sSSY_LOWER_2+", and "+annual_sSSY_TOTAL_2+" tons/km"u"\u00B2""/year, respectively.")


## Annual SSY and sSSY from Table 4 (UPPER, LOWER_QUARRY, LOWER_VILLAGE, LOWER)
document.add_paragraph("For storms with measured SSY at FG1, FG2, and FG3 (Table "+S_budget_2_table.table_num+") Psmeas was "+P_measured_4+" mm, or "+"%.0f"%P_measured_4_perc_storm+"% of Psann. Using Equation "+SSY_annual_eq.eq_num+", estimated annual SSY from the UPPER, LOWER_QUARRY, LOWER_VILLAGE, LOWER, and TOTAL watersheds was "+annual_SSY_UPPER_4+", "+annual_SSY_LOWER_QUARRY_4+", "+annual_SSY_LOWER_VILLAGE_4+", "+annual_SSY_LOWER_4+", and "+annual_SSY_TOTAL_4+" tons/year, respectively. Annual sSSY from the UPPER, LOWER_QUARRY, LOWER_VILLAGE, LOWER, and TOTAL watersheds were estimated to be "+annual_sSSY_UPPER_4+", "+annual_sSSY_LOWER_QUARRY_4+", "+annual_sSSY_LOWER_VILLAGE_4+", "+annual_sSSY_LOWER_4+" and "+annual_sSSY_TOTAL_4+" tons/km"u"\u00B2""/year, respectively.")

### DISCUSSION
document.add_heading('Discussion',level=2)

## Methods for quantifying human impact
document.add_heading('Methods for quantifying human impact',level=3)

# What have we learned about how humans impact sediment loads on remote tropical islands?  How does your SSY compare with other mined areas?  About the best predictors of SSY?  What was the utility of turbidity measuremetns, and have others constructed sediment budgets like this with tubidity?  What was their cumulative RMSE?  You have much of this in the results, but I think the impact of the paper will be stronger if you interpret the results vis-à-vis other work in a separate discussion.

document.add_paragraph("In contrast to other methods like USLE-based models or traditional sediment rating curves (SSC-Q), event-wise correlation of SSYEV and storm metrics was advantageous for quantifying increased sediment loading from human-disturbed areas in the study watershed. USLE-based models are not well-calibrated for use in steep, tropical watersheds with human disturbance (Calhoun and Fletcher, 1999; Sadeghi et al., 2007), and have high uncertainty in the sediment delivery ratio. Using a traditional relationship between Q and SSC to estimate continuous sediment load was problematic in the study watershed, due to the significant hysteresis and changing sediment availability. While the Q-SSC relationship illustrated large differences in SSC downstream of the quarry and reduced SSC at low Q following mitigation activities, the Q-SSC method would not provide accurate estimates of sediment loading to the coast and coral reef ecosystem.")

document.add_paragraph("Reid and Dunne (1996) argue that in cases where there is a clear management question and the study area can be divided into sub-units, a sediment budget can be rapidly developed with only a few field measurements and limited periods of field monitoring. In Faga'alu watershed, and other similar steep watersheds, human-disturbance is often constrained to the lower watershed, and sediment yields from these key sources can be measured separately from the undisturbed forest upstream. Sampling in Faga'alu watershed targeted key sediment sources, and the disturbance signal was very large. Analyzing  event-wise SSY allows comparison of similar size storms to determine change over time without problems of interannual variability in precipitation totals, and eliminates the need for long-term continuous field work to measure annual total yields. From a management perspective, this approach is less expensive since it does not require multiple or even a single full year of monitoring, and can be rapidly conducted if mitigation or disturbance activities are already planned. By developing a predictive model of SSYEV from an easily monitored storm metric like maximum event discharge, SSYEV can be modeled in the future to compare with either post-mitigation or post-disturbance SSYEV.")

## Assessing alpha and Beta model parameters
document.add_heading('Interpreting slope and intercept of the Qmax-SSY relationship',level=3)
All_Models_stats
alpha_upper, beta_upper = "%.2f"%float(All_Models_stats['alpha'].loc['Qmax_upper']), "%.2f"%float(All_Models_stats['Beta'].loc['Qmax_upper'])
alpha_total, beta_total = "%.1f"%float(All_Models_stats['alpha'].loc['Qmax_total']), "%.1f"%float(All_Models_stats['Beta'].loc['Qmax_total'])
alpha = u"\u03B1"
beta = u"\u03B2"

document.add_paragraph("Several researchers have attempted to explain values of the intercept ("u"\u03B1"") and slope ("u"\u03B2"") coefficients of the sediment rating curve as a function of watershed characteristics. A traditional sediment rating curve (Q-SSC) is considered a 'black box' model, and though the slope and intercept have no physical meaning, some physical interpretation has been ascribed to them (Asselman, 2000). Rankl (2004) hypothesized that the intercept in the Qmax-SSYEV relationship varied with sediment availability and erodibility in watersheds. Duvert et al. (2012) found that intercepts of the Qmax-SSYEV relationship are also dependent on the regression fitting method. While slopes in log-log space can be compared directly (Duvert et al., 2012), intercepts must be plotted in similar units, and similarly normalized by watershed area. In five semi-arid to arid watersheds (2.1-1,538 km"u"\u00B2"") in Wyoming, United States, Rankl (2004) found intercepts of the SSYEV-Qmax relationship ranged from 111-4,320 (Qmax in m"u"\u00B3""/s/km"u"\u00B2"", SSY in Mg/km"u"\u00B2""). In eight sub-humid to semi-arid watersheds (0.45-22 km"u"\u00B2""), Duvert et al. (2012)  found the intercepts ranged from 25-5,039. In Faga'alu, the intercept in the undisturbed, UPPER subwatershed was 0.35, and in the disturbed, TOTAL watershed the intercept was 1.38, which are an order of magnitude or two lower than the lowest intercepts in Rankl (2004) and Duvert et al. (2012). This suggests that sediment availability is relatively low in Faga'alu, under natural and human-disturbed conditions, likely due to the dense forest cover.")

document.add_paragraph("High slope values  in the log-log plots ("u"\u03B2"" coefficient) suggest that small changes in stream discharge lead to large increases in sediment load due to the erosive power of the river or the availability of new sediment sources at high Q (Asselman, 2000). Rankl (2004) assumed that the slope was a function of rainfall intensity on hillslopes, and found that the slopes ranged from 1.07-1.29 in five semi-arid to arid watersheds in Wyoming, and were not statistically different among watersheds. In the watersheds in Duvert et al. (2012), slopes ranged from 0.95-1.82, and from 1.06-2.45 in eighteen other watersheds (0.60-1,538 km"u"\u00B2"") in diverse geographical settings (Basher et al., 1997; Fahey and Marden, 2000; Hicks et al., 2009; Rankl, 2004; Tropeano, 1991) compiled by Duvert et al. (2012). In Faga'alu, slopes were "+beta_upper+" and "+beta_total+" in the UPPER and TOTAL Faga'alu watersheds, respectively, which are very consistent with the watersheds presented in Rankl (2004) and Duvert et al. (2012).")


document.add_paragraph("In Faga'alu, SSYEV was least correlated with the Erosivity Index (EI30). Duvert et al. (2012) also found low correlation coefficients with 5 min rainfall intensity for 8 watersheds in France and Mexico. Rodrigues et al. (2013) hypothesized that EI30 is poorly correlated with SSY due to the effect of previous events on antecedent moisture conditions and in-channel sediment storage. Cox et al. (2006) found EI30 was more correlated with soil loss in an agricultural watershed than a forested watershed, and Faga'alu is mainly covered in dense forest. Similar to other studies (Basher et al., 2011; Duvert et al., 2012; Fahey et al., 2003; Hicks, 1990; Rankl, 2004; Rodrigues et al., 2013) the highest correlations with SSYEV at Faga'alu were observed for discharge metrics, particularly Qmax which had the highest correlation of the tested storm metrics.")

## compare sSSY and SSC with other watersheds
## How does the yield from the quarry differ from other studies (e.g. Wolman 1967, but also look up others)?  

document.add_heading('Comparing sSSY and SSC in other small Pacific Island watersheds',level=3)

# regional sediment yields
document.add_paragraph("Sediment yield is highly variable among individual watersheds, but is generally controlled by climate, vegetation cover, and geology, with human disturbance playing an increasing role in the 20th century (Syvitski et al., 2005). Sediment yields in tropical Southeast Asia and high-standing islands between Asia and Australia range from ~10 tons/km"u"\u00B2""/yr in the granitic Malaysian Peninsula to ~10,000 tons/km"u"\u00B2""/yr in the tectonically active, steeply sloped island of Papua New Guinea (Douglas, 1996). Sediment yields from Faga'alu are on the lower end of the range, with sSSY of 33-80 tons/km"u"\u00B2""/yr from the undisturbed UPPER watershed, and 170-380 tons/km"u"\u00B2""/yr from the disturbed TOTAL watershed.")

# compare to Milliman's models
document.add_paragraph("Milliman and Syvitski (1992) report there is an unusually high average sSSY of 1,000-3,000 tons/km"u"\u00B2""/year, from watersheds (10-100,000 km"u"\u00B2"") in tropical Asia and Oceania. However, Milliman and Syvitski's (1992) regional models of sSSY as a function of basin size and maximum elevation predict only 13 tons/km"u"\u00B2""/year from watersheds with peak elevation 500-1,000 m (highest point of UPPER Faga'alu subwatershed is 653 m), but 68 tons/km"u"\u00B2""/year for max elevations of 1,000-3,000 m,   which is comparable to SSY measured from UPPER Faga'alu watershed (33-80 tons/km"u"\u00B2""/yr) (Table 8). Given the high vegetation cover and lack of human activity in the UPPER Faga'alu subwatershed, it is assumed that sSSY should be several orders of magnitude lower than watersheds presented in Milliman and Syvitski (1992)   but sSSY from the forested UPPER Faga'alu subwatershed was approximately two times higher. However, the UPPER subwatershed is a smaller watershed than included in Milliman and Syvitski (1992) (smallest 100 km"u"\u00B2""), with less sediment storage, and orders of magnitude scatter around their model is observed for smaller watersheds (Figures 5e and 6e in Milliman and Syvitski 1992).")

# compare to Hawaii watersheds
document.add_paragraph("Few examples of sediment yield studies on volcanic, Pacific Islands similar to Tutuila were found in the literature for comparison, except for studies in two Hawaiian watersheds (Stock and Tribble, 2010): Hanalei watershed on Kauai, and Kawela watershed on Molokai (Table "+lit_values_for_Annual_ssy_table_num+"). Where mean and maximum SSC values were similar to Faga'alu, sSSY was also similar. Hanalei watershed on Kauai (54 km"u"\u00B2""), has similarly steep relief and high rainfall (varies with elevation from 2,000-11,000 mm), with average SSC of 63 mg/L and maximum SSC of 2,750 mg/l. Calhoun and Fletcher (1999) previously estimated sSSY from Hanalei was 140"u"\u00B1""55 tons/km"u"\u00B2""/year, but had fewer data than Stock and Tribble (2010), who estimated sSSY was 525 tons/km"u"\u00B2""/yr. In Kawela watershed on Molokai (14 km"u"\u00B2""), a grazing-disturbed sub-humid watershed (precipitation varies with elevation from 500-3,000 mm), Stock and Tribble (2010) estimated sSSY was 459 tons/km"u"\u00B2""/yr. In Kawela, Molokai, SSC was much higher than measured in Faga'alu and occurred at relatively lower flow, with average SSC of 3,490 mg/L, and a maximum value of 54,000 mg/L at an instantaneous flow of 1.614 m"u"\u00B3""/sec. In the disturbed, TOTAL Faga'alu watershed where average SSC was  "+Mean_SSC_FG3+" mg/L and maximum SSC was "+Max_SSC_FG3+" mg/L, annual sSSY was estimated to be 214-250 tons/km"u"\u00B2""/yr, similar to Hanalei watershed.")


# compare specific land covers
document.add_paragraph("Annual sSSY from the disturbed quarry was estimated to be approximately "+annual_sSSY_disturbed_LOWER_QUARRY_4+" tons/km"u"\u00B2""/yr. The quarry surfaces are comprised of haul roads, piles of overburden, and steep rock faces which can be described as a mix of unpaved roads and cut-slopes. Literature values show measured sSSY from cutslopes varying from 0.01 tons/km"u"\u00B2""/yr in Idaho (Megahan, 1980) to 105,000 tons/km"u"\u00B2""/yr in Papua New Guinea (Blong and Humphreys, 1982), so the sSSY ranges measured in this study are well within the ranges found in the literature.")


## compare with other kinds of disturbance
document.add_heading('Comparison with other kinds of sediment disturbance', level=3)

## Is the 3.8x increase big or small compared with others?
document.add_paragraph("Other studies in small, mountainous watersheds have documented one to several orders of magnitude increases in SSY from land use that disturbs a small fraction of the watershed area. Urbanization and mining increase sediment yield in stable terrain by two to three orders of magnitudes in catchments of several km"u"\u00B2"". Yields from construction sites can exceed those from the most unstable, tectonically active natural environments of Southeast Asia (Douglas, 1996). In Kawela watershed on Molokai, less than 5% of the land produces most of the sediment, and only 1% produces ~50% of the sediment (Risk, 2014; Stock et al., 2010)   . In three basins on St. John, US Virgin Islands, Ramos-Scharr"u"\u00F3""n and Macdonald (2005) found unpaved roads increased sediment delivery rates by 3-9 times. Disturbances at larger scales have had similar increases in total SSY to coral environments. The development of the Great Barrier Reef (GBR) catchment since European settlement (ca.1830) led to increases in SSY by an estimated factor of 5.5 x (Kroon et al., 2012). Mining activity has been a major contributor of sediment in other watersheds on volcanic islands with steep topography and high precipitation, i ncreasing sediment yields by 5-10 times in a watershed in Papua New Guinea (Hettler et al., 1997; Thomas et al., 2003). In contrast to other land disturbances like fire, logging, or urbanization where sediment disturbance decreases over time, the disturbance from mining is persistently high. Disturbance magnitudes are similar to the construction phase of urbanization (Wolman and Schick, 1967), or high-traffic unpaved roads (Reid and Dunne, 1984), but persist or even increase over time.")

#### CONCLUSION
conclusion_title=document.add_heading('Conclusion',level=2)

document.add_paragraph("Human disturbance has increased sediment yield to Faga'alu Bay by "+S_budget['TOTAL tons']['DR']+"x over pre-disturbance levels. The human-disturbed subwatershed accounted for the majority ("+S_budget['% LOWER']['Total/Avg']+"%) of total sediment yield, and the quarry (1.1% of watershed area) contributed almost half of total SSY to the Bay. Qmax was the best predictor of SSYEV. The slopes of the Qmax-SSYEV relationships were comparable with other studies, but the model intercepts were an order of magnitude lower than intercepts from watersheds in semi-arid to semi-humid climates. This suggests that sediment availability is relatively low in the Faga'alu watershed, either because of the heavy forest cover or volcanic rock type. The relative contribution from the human-disturbed watershed was hypothesized to diminish with increasing storm size but the results from precipitation metrics and discharge metrics were contradictory. The Psum-SSYEV model showed that the relative contribution of SSYEV from the human-disturbed watershed decreases with storm size, but the Qmax-SSYEV model shows no change in relative contributions over increasing storm size. It was hypothesized that SSY from natural areas would become the dominant source for larger storm events, but the DR remains high for large storm events due to the naturally low SSY from natural forest areas in Faga'alu watershed.  This suggests that disturbed areas were not supply limited for the range of sampled storms.")

## Introduce the post-mitigation work
document.add_paragraph("Management has responded to data on sediment loading in Faga'alu and has undertaken a sediment mitigation program. In August 2012, preliminary results of the significant SSYEV contributions from the quarry and its impact on coral reef health in Faga'alu Bay were communicated to US Federal and local environmental management and conservation groups including the Faga'alu village community, NOAA Coral Reef Conservation Program, American Samoa Environmental Protection Agency, and the American Samoa Coral Reef Advisory Group. In February 2013, Faga'alu watershed was designated by the US Coral Reef Task Force as a Priority Watershed Restoration site, with the main objective to reduce sediment yields to the adjacent coral reefs. These groups developed a sediment management plan for the quarry operators and village residents which was implemented in October 1, 2014, and completed in December 2014. Storm monitoring is currently in progress and results documenting the successful reduction of sediment yields to Faga'alu Bay will be presented in a forthcoming paper. For a full description of the sediment mitigation project and documentation, see Holst-Rice et al. (2015). This work provides an example of an environmental management project which could only be accomplished by the effective partnerships between community groups, local industries, educational institutions, and government regulatory and funding agencies.")

#### Acknowledgements
document.add_heading('Acknowledgements', level=2)
document.add_paragraph("Funding for this project was provided by NOAA Coral Reef Conservation Program (CRCP) through the American Samoa Coral Reef Advisory Group (CRAG). Kristine Bucchianeri at CRAG and Susie Holst at NOAA CRCP provided necessary and significant support. Christianera Tuitele, Phil Wiles (currently at the South Pacific Regional Environment Programme), and Tim Bodell at American Samoa Environmental Protection Agency (ASEPA), and Fatima Sauafea-Leau and Hideyo Hattori at NOAA, provided on-island coordination with traditional local authorities. Dr. Mike Favazza provided critical logistical assistance in American Samoa. Robert Koch at the American Samoa Coastal Zone Management Program (ASCMP) and Travis Bock at ASEPA assisted in accessing historical geospatial and water quality data. Many others helped and supported the field and laboratory work including Professor Jameson Newtson, Rocco Tinitali, and Valentine Vaeoso at American Samoa Community College (ASCC), Meagan Curtis and Domingo Ochavillo at American Samoa Department of Marine and Wildlife Resources (DMWR), Don and Agnes Vargo at American Samoa Land Grant, Christina Hammock at NOAA American Samoa Climate Observatory, and Greg McCormick at San Diego State University. George Poysky, Jr., George Poysky III, and Mitch Shimisaki at Samoa Maritime Ltd. provided unrestricted access to the Faga'alu quarry site, and historical operation information. Faafetai tele lava.")

#### REFERENCES
document.add_heading('References', level=2)

#### Appendix 1
document.add_page_break()
document.add_heading('APPENDIX 1. Channel cross sections',level=2)

if 'LBJ_Cross_Section' in locals():
    document.add_picture(LBJ_Cross_Section['filename']+'.png',width=Inches(6))
    add_figure_caption(LBJ_Cross_Section['fig_num'],"Stream cross-section at FG3")

if 'DAM_Cross_Section' in locals():  
    document.add_picture(DAM_Cross_Section['filename']+'.png',width=Inches(6))
    add_figure_caption(DAM_Cross_Section['fig_num'],"Stream cross-section at FG1")
    
#### Appendix 2
document.add_page_break()
document.add_heading("APPENDIX 2. Dams in Faga'alu watershed",level=2)

## Fagaalu Reservoir Infrastructure
document.add_paragraph("Faga'alu stream was dammed at 4 locations above the village: 1) Matafao Dam (elevation 244 m) near the base of Mt. Matafao, draining 0.20 km"u"\u00B2"", 2) Vaitanoa Dam at Virgin Falls (elevation 140 m), draining an additional 0.44 km"u"\u00B2"", 3) a small unnamed dam below Vaitanoa Dam at elevation 100m, and 4) Lower Faga'alu Dam (elevation 48 m), immediately upstream of a large waterfall 30 m upstream of the quarry, draining an additional 0.26 km"u"\u00B2"" (Tonkin & Taylor International Ltd. 1989). A 2012 aerial LiDAR survey (Photo Science, Inc.) indicates the drainage area at the Lower Faga'alu Dam is 0.90 km"u"\u00B2"". A small stream capture/reservoir (~35 m"u"\u00B3"") is also present on a side tributary that joins Faga'alu stream on the south bank, opposite the quarry. It is connected to a ~6 cm diameter pipe but it is unknown when or by whom it was built, its initial capacity, or if it is still conveying water. During all site visits water was overtopping this small structure through the spillway crest, suggesting it is fed by a perennial stream.")

document.add_paragraph("Matafao Dam was constructed in 1917 for water supply to the Pago Pago Navy base, impounding a reservoir with initial capacity of 1.7 million gallons (6,400 m"u"\u00B3"") and piping the flow out of the watershed to a hydropower and water filtration plant in Fagatogo. In the early 1940's the Navy replaced the original cement tube pipeline and hydropower house with cast iron pipe but it is unknown when the scheme fell out of use (Tonkin & Taylor International Ltd. 1989; URS Company 1978). Remote sensing and a site visit on 6/21/13 confirmed the reservoir is still filling to the spillway crest with water and routing some flow to the Fagatogo site, though the amount is much less than the 10 in. diameter pipes conveyance capacity and the flow rate variability is unknown. A previous site visit on 2/21/13 by American Samoa Power Authority (ASPA) found the reservoir empty of water but filled with an estimated 3-5 meters of fine sediment (Kearns 2013). Interviews with local maintenance staff and historical photos confirmed the Matafao Reservoir was actively maintained and cleaned of sediment until the early 70's.")

document.add_paragraph("The Vaitanoa (Virgin Falls) Dam, was built in 1964 to provide drinking water but the pipe was not completed as of 10/19/89, and a stockpile of some 40 (8 ft length) 8 in. diameter asbestos-cement pipes was found on the streambanks. Local quarry staff recall the pipes were removed from the site some time in the 1990's. The Vaitanoa Reservoir had a design volume of 4.5 million gallons (17,000m"u"\u00B3""), but is assumed to be full of sediment since the drainage valves were never opened and the reservoir was overtopping the spillway as of 10/18/89 (Tonkin & Taylor International Ltd. 1989). A low masonry weir was also constructed downstream of the Vaitanoa Dam, but not connected to any piping.") 

document.add_paragraph("The Lower Faga'alu Dam was constructed in 1966/67 just above the Samoa Maritime, Ltd. Quarry, as a source of water for the LBJ Medical Centre. It is unknown when this dam went out of use but in 1989 the 8 in. conveyance pipe was badly leaking and presumed out of service. The 8 in. pipe disappears below the floor of the Samoa Maritime quarry and it is unknown if it is still conveying water or has plugged with sediment. The derelict filtration plant at the entrance to the quarry was disconnected prior to 1989 (Tonkin & Taylor International Ltd. 1989). The original capacity was 0.03 million gallons (114 m"u"\u00B3"") but is now full of coarse sediment up to the spillway crest. No reports were found indicating this structure was ever emptied of sediment.") 

#### Appendix 3
document.add_page_break()
document.add_heading("APPENDIX 3. Water discharge during storm events",level=2)
#document.add_paragraph("The DR for specific water discharge (Q) was also calculated to determine if observed changes in SSY were attributable to errors in quantifying Q, and if urbanization and agriculture has affected the total water discharge from the watershed. The DR for water discharge was "+DR_Q+", suggesting that specific water discharge from the subwatersheds is the same, which is expected since they are similar sizes. Urbanization has likely increased water discharge, since the relatively large amounts of impervious surface in the village area increase runoff. However, the village area occupies a small percentage ("+"%.1f"%landcover_table.ix[2]['% High Intensity Developed']+"%) of the LOWER subwatershed area in comparison to the forest area ("+"%.1f"%landcover_table.ix[2]['% Forest']+") (Table "+landcover_table.table_num+"), and the relatively small increase in Q could be within the acceptable measurement error for quantifying Q.")  

## Storm Water Discharge
if 'Q_Diff_table' in locals():
    dataframe_to_table(df=Q_Diff_table,table_num=Q_Diff_table.table_num,caption="Water discharge from subwatersheds in Faga'alu",fontsize=9)
## 

#### Appendix 4
document.add_page_break()
document.add_heading("APPENDIX 4. Synthetic rating curves for turbidimeters in Faga'alu",level=2)

## Synthetic Rating Curves Fagaalu
if 'Synthetic_Rating_Curve' in locals():
    document.add_picture(Synthetic_Rating_Curve['filename']+'.png',width=Inches(6)) ## add pic from filename defined above
    add_figure_caption(Synthetic_Rating_Curve['fig_num'],"Synthetic Rating Curves for (a) OBS turbidimeter deployed at FG3 and (b) YSI deployed at FG1.")

## Special Unicode letters
#m3 = 'm'u"\u00B3"
#km2 = 'km'u"\u00B2"' other text'
#SSYEV = 'SSY'u"\u2091"u"\u1D65"
#alpha = u"\u03B1"
#beta = u"\u03B2"
#o_accent = ''
#e_accent = '' 
#plus_minus = u"\u00B1"

#upper_code = '<w:r><w:rPr><vertAlign w:val="superscript"/></w:rPr><w:t>UPPER</w:t></w:r>'
#
#for paragraph in document.paragraphs:
#    paragraph
#    ## superscript these
#    paragraph.text = paragraph.text.replace('km2','km'u"\u00B2")
#    paragraph.text = paragraph.text.replace('m3','m'u"\u00B3")
#    paragraph.text = paragraph.text.replace('r2','r'u"\u00B2")
#    paragraph.text = paragraph.text.replace('SSYUPPER','SSY' +upper_code)


   
## Save Document
document.save(maindir+'Manuscript/DRAFT-Fagaalu_Sediment_Yield_2015.docx')


## Clean up any open figures
plt.close('all')








