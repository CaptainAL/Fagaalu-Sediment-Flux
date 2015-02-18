# -*- coding: utf-8 -*-
"""
Created on Tue Feb 17 10:50:00 2015

@author: Alex
"""
plt.ioff()
from docx import *
from docx.shared import Inches
##  Create Document
document = Document()
## Add  Heading 
document.add_heading('Turbidity to SSC',level=0)
## Add  paragraph text
paragraph = document.add_paragraph('This section describes the turbidity to SSC process:')
## Add  table
document.add_paragraph('This is a table')
table = document.add_table(rows=0, cols=3)
tbl_rws = [['A1','B1','C1'],
           ['A2','B2','C2'],
           ['A3','B3','C3']]
for row in tbl_rws:
    row_cells = table.add_row().cells
    row_cells[0].text = row[0]
    row_cells[1].text = row[1]
    row_cells[2].text = row[2]
table.style = 'TableGrid'    

document.add_heading('Synthetic Rating Curves for Turbidimeters',level=1)
## Add Figure
fig1_filename = figdir+'Synthetic Rating Curves.png' ## define file name to find the png file from other script
## generate figure from separate script and save to file
fig1 = Synthetic_Rating_Curves(param='SS_Mean',show=False,save=True,filename=fig1_filename)
document.add_picture(fig1_filename,width=Inches(7)) ## add pic from filename defined above

document.save('demo.docx')










